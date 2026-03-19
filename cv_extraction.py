# ─────────────────────────────────────────────────────────────────────────────
# cv_extraction.py — CV text and photo extraction
# Imported by main.py:  from cv_extraction import (...)
# ─────────────────────────────────────────────────────────────────────────────
import os
import re
import io
import base64
import zipfile
import struct
import logging
import os
from openai import AzureOpenAI
 
log = logging.getLogger("cv_booster")
 
# Azure Vision client (utilise TES variables .env)
AZURE_VISION_DEP = (
    os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME") or
    os.getenv("AZURE_OPENAI_DEPLOYMENT", "gpt-4o-mini")
)
 
client = AzureOpenAI(
    api_version    = os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview"),
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT", ""),
    api_key        = os.getenv("AZURE_OPENAI_API_KEY", ""),
)
try:
    import fitz
except ImportError:
    fitz = None

import pdfplumber
import docx as docx_reader


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1 — TEXT EXTRACTION
# ═════════════════════════════════════════════════════════════════════════════

def _is_thin(text: str) -> bool:
    """True when extracted text is too short / garbled to be a complete CV."""
    if not text or len(text) < 120:
        return True
    lines = [l for l in text.split("\n") if l.strip()]
    return len(lines) < 8 or len(text.split()) < 30


# ── PDF → base64 images (for GPT-4o vision) ──────────────────────────────────

def _pdf_to_images_b64(file_bytes: bytes, max_pages: int = 4) -> list:
    """Render each PDF page to PNG at 150 DPI, return list of base64 strings."""
    if not FITZ_OK:
        return []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        out = []
        for i in range(min(len(doc), max_pages)):
            mat = fitz.Matrix(150 / 72, 150 / 72)
            pix = doc[i].get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            out.append(base64.b64encode(pix.tobytes("png")).decode())
        return out
    except Exception as e:
        log.error("_pdf_to_images_b64 error: %s", e)
        return []


# ── DOCX → PDF via LibreOffice → base64 images ───────────────────────────────

def _docx_to_images_b64(file_bytes: bytes) -> list:
    """Convert DOCX to PDF via LibreOffice headless, then render pages."""
    import subprocess, tempfile
    try:
        with tempfile.TemporaryDirectory() as tmp:
            p_in = os.path.join(tmp, "cv.docx")
            with open(p_in, "wb") as f:
                f.write(file_bytes)
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, p_in],
                capture_output=True, timeout=30
            )
            p_out = os.path.join(tmp, "cv.pdf")
            if result.returncode != 0 or not os.path.exists(p_out):
                log.error("LibreOffice failed: %s", result.stderr.decode()[:200])
                return []
            return _pdf_to_images_b64(open(p_out, "rb").read())
    except Exception as e:
        log.error("_docx_to_images_b64 error: %s", e)
        return []


# ── GPT-4o Vision extraction ──────────────────────────────────────────────────

def _vision_extract(images_b64: list) -> str:
    """
    Send CV page images to GPT-4o vision.
    Handles ALL graphical layouts:
      - Two-column / sidebar CVs
      - Skill bars → converts to text level
      - Star / dot ratings → converts to level
      - Language flags → extracts name + level
      - Tables, text boxes, shapes
      - Scanned PDFs (OCR via vision)
    """
    if not images_b64:
        return ""

    content = [
        {
            "type": "text",
            "text": (
                "You are a CV text extractor. Extract ALL text completely — do not skip or truncate anything.\n\n"
                "LAYOUT RULES:\n"
                "- Two-column / sidebar: extract LEFT column top→bottom COMPLETELY, then RIGHT column top→bottom COMPLETELY\n"
                "- Preserve ALL section headers exactly as written\n"
                "- Preserve ALL bullet points — start each with '- '\n"
                "- Preserve ALL projects — do not skip any project\n"
                "- Preserve dates and locations exactly\n"
                "- Extract ALL certifications listed\n\n"
                "SKILL BARS: estimate fill → ~20%=Beginner ~60%=Intermediate ~80%=Advanced ~95%=Expert\n"
                "STAR RATINGS: 1-2/5=Beginner 3/5=Intermediate 4/5=Advanced 5/5=Expert\n"
                "LANGUAGE FLAGS: extract language name and CEFR level if shown\n"
                "ICONS (phone, email, location, LinkedIn): skip icon, extract text only\n"
                "PHOTOS: skip entirely\n"
                "OUTPUT: plain text only, no markdown, no asterisks, no commentary.\n"
                "Separate pages with a blank line.\n"
                "CRITICAL: Extract every single item — projects, certifications, experiences — completely. Do not stop early."
            )
        }
    ]
    for b64 in images_b64:
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{b64}", "detail": "high"}
        })

    try:
        resp = client.chat.completions.create(
            model=AZURE_VISION_DEP,
            messages=[{"role": "user", "content": content}],
            temperature=0,
            max_tokens=3000,
        )
        result = resp.choices[0].message.content.strip()
        log.info("✓ Vision extraction: %d chars", len(result))
        return result
    except Exception as e:
        log.error("Vision extraction failed: %s", e)
        return ""


# ── XML direct extraction for DOCX ───────────────────────────────────────────

def _find_xml_parent(root, child):
    for parent in root.iter():
        if child in list(parent):
            return parent
    return None


def _xml_extract_docx(file_bytes: bytes) -> str:
    """
    Parse DOCX XML directly.
    - Removes mc:Fallback nodes → eliminates AlternateContent duplicates
    - Collects ALL w:p paragraph text in document order
    - Deduplicates with ordered set
    """
    import xml.etree.ElementTree as ET
    W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            xml_bytes = z.read("word/document.xml")
    except Exception as e:
        raise RuntimeError(f"Cannot open DOCX zip: {e}")

    root = ET.fromstring(xml_bytes)
    for fb in list(root.iter(f"{{{MC}}}Fallback")):
        parent = _find_xml_parent(root, fb)
        if parent is not None:
            parent.remove(fb)

    seen, lines = set(), []
    for p_el in root.iter(f"{{{W}}}p"):
        parts = [t.text for t in p_el.iter(f"{{{W}}}t") if t.text]
        text  = "".join(parts).strip()
        if text and text not in seen:
            seen.add(text); lines.append(text)

    return "\n".join(lines)


# ── Main extraction functions ─────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    PDF extraction with 2-level fallback:
      1. pdfplumber smart — detects true 2-column layouts via X gap analysis
      2. GPT-4o vision — handles graphical, sidebar, scanned PDFs
    """
    best = ""

    try:
        all_lines = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                words = page.extract_words(x_tolerance=3, y_tolerance=3)
                if not words:
                    continue

                page_width = page.width

                # ── Detect TRUE 2-column layout via gap analysis ──────────────
                # A real 2-col layout has a clear empty gap zone in the middle X range
                mid_lo = page_width * 0.35
                mid_hi = page_width * 0.55
                mid_words = [w for w in words if mid_lo < w["x0"] < mid_hi]
                left_words  = [w for w in words if w["x0"] <= mid_lo]
                right_words = [w for w in words if w["x0"] >= mid_hi]

                is_two_col = (
                    len(mid_words) < len(words) * 0.08 and  # less than 8% words in middle gap
                    len(left_words) > 15 and
                    len(right_words) > 15
                )

                if is_two_col:
                    log.info("✓ PDF: 2-column layout detected (gap: %d%% words in middle)",
                             int(len(mid_words)/len(words)*100))
                    # Extract left column top→bottom
                    left_lines = {}
                    for w in left_words:
                        y = round(w["top"] / 4) * 4
                        left_lines.setdefault(y, []).append(w["text"])
                    # Extract right column top→bottom
                    right_lines = {}
                    for w in right_words:
                        y = round(w["top"] / 4) * 4
                        right_lines.setdefault(y, []).append(w["text"])

                    for y in sorted(left_lines):
                        all_lines.append(" ".join(left_lines[y]))
                    all_lines.append("")
                    for y in sorted(right_lines):
                        all_lines.append(" ".join(right_lines[y]))
                else:
                    # Single-column page — but local 2-column grids (projects, skills,
                    # certs) may exist.  Detect per-line gaps >= MIN_COL_GAP and split
                    # left/right into separate blocks, flushed together to preserve order.
                    MIN_COL_GAP = 60  # px — well above normal word spacing (2-8px)

                    lines_map = {}
                    for w in words:
                        y = round(w["top"] / 4) * 4
                        lines_map.setdefault(y, []).append(w)

                    left_block:  list = []
                    right_block: list = []

                    def _flush_blocks():
                        for _, lt in sorted(left_block):
                            all_lines.append(lt)
                        for _, rt in sorted(right_block):
                            all_lines.append(rt)
                        left_block.clear()
                        right_block.clear()

                    for y in sorted(lines_map):
                        row_words = sorted(lines_map[y], key=lambda w: w["x0"])
                        split_idx = None
                        for i in range(1, len(row_words)):
                            if row_words[i]["x0"] - row_words[i - 1]["x1"] >= MIN_COL_GAP:
                                split_idx = i
                                break

                        if split_idx is not None:
                            lt = " ".join(w["text"] for w in row_words[:split_idx])
                            rt = " ".join(w["text"] for w in row_words[split_idx:])
                            _r = rt.strip()
                            # Skip split for short metadata labels, dates, or degree types
                            _is_short  = len(_r.split()) <= 2
                            _is_date   = bool(re.search(
                                r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec"
                                r"|janv|févr|mars|avr|juin|juil|août|sept|déc)\b",
                                _r, re.I))
                            _is_degree = bool(re.match(
                                r"^(master|bachelor|licence|diplôme|ingénieur|doctorat"
                                r"|bts|dut|bac|mba|classes\s+prép)",
                                _r, re.I))
                            if _is_short or _is_date or _is_degree:
                                _flush_blocks()
                                all_lines.append(" ".join(w["text"] for w in row_words))
                            else:
                                left_block.append((y, lt))
                                right_block.append((y, rt))
                        else:
                            _flush_blocks()
                            all_lines.append(" ".join(w["text"] for w in row_words))

                    _flush_blocks()

        text = "\n".join(all_lines)
        if not _is_thin(text):
            log.info("✓ PDF: pdfplumber OK (%d chars)", len(text))
            return text
        if len(text) > len(best):
            best = text
        log.info("⚡ PDF: pdfplumber thin (%d chars) → vision", len(text))
    except Exception as e:
        log.error("pdfplumber error: %s", e)

    # Level 2: GPT-4o vision
    imgs = _pdf_to_images_b64(file_bytes)
    if imgs:
        vt = _vision_extract(imgs)
        if vt and len(vt) > 60:
            log.info("✓ PDF: vision OK (%d chars)", len(vt))
            return vt

    log.warning("⚠  PDF: using partial result (%d chars)", len(best))
    return best


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    DOCX extraction with 3-level fallback:
      1. XML direct  — best for templates with mc:AlternateContent, text boxes
      2. python-docx — paragraphs + all table cells
      3. LibreOffice + GPT-4o vision — for complex graphical DOCX
    """
    best = ""

    # Level 1: XML direct
    try:
        t = _xml_extract_docx(file_bytes)
        if not _is_thin(t):
            log.info("✓ DOCX: XML direct OK (%d chars)", len(t))
            return t
        if len(t) > len(best): best = t
        log.info("⚡ DOCX: XML thin (%d chars) → python-docx", len(t))
    except Exception as e:
        log.error("XML extract error: %s", e)

    # Level 2: python-docx (paragraphs + tables)
    try:
        doc  = docx_reader.Document(io.BytesIO(file_bytes))
        seen, lines = set(), []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t and t not in seen: seen.add(t); lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = p.text.strip()
                        if t and t not in seen: seen.add(t); lines.append(t)
        t = "\n".join(lines)
        if not _is_thin(t):
            log.info("✓ DOCX: python-docx OK (%d chars)", len(t))
            return t
        if len(t) > len(best): best = t
        log.info("⚡ DOCX: python-docx thin (%d chars) → vision", len(t))
    except Exception as e:
        log.error("python-docx error: %s", e)

    # Level 3: LibreOffice → PDF → GPT-4o vision
    imgs = _docx_to_images_b64(file_bytes)
    if imgs:
        vt = _vision_extract(imgs)
        if vt and len(vt) > 60:
            log.info("✓ DOCX: LibreOffice+vision OK (%d chars)", len(vt))
            return vt

    log.warning("⚠  DOCX: using best partial result (%d chars)", len(best))
    return best


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PHOTO EXTRACTION
# ═════════════════════════════════════════════════════════════════════════════

def _img_dimensions(img_bytes: bytes, ext: str):
    try:
        if ext == ".png":
            w = struct.unpack(">I", img_bytes[16:20])[0]
            h = struct.unpack(">I", img_bytes[20:24])[0]
            return w, h
        if ext in (".jpg", ".jpeg"):
            i = 0
            while i < len(img_bytes) - 4:
                if img_bytes[i] == 0xFF and img_bytes[i+1] in (0xC0, 0xC1, 0xC2):
                    h = struct.unpack(">H", img_bytes[i+5:i+7])[0]
                    w = struct.unpack(">H", img_bytes[i+7:i+9])[0]
                    return w, h
                i += 1
    except Exception:
        pass
    return 100, 100


def _is_portrait(w: int, h: int) -> bool:
    if h == 0: return False
    ratio = w / h
    return 0.5 < ratio < 1.4 and 60 <= w <= 800


def extract_photo_from_pdf(file_bytes: bytes):
    if not FITZ_OK: return None
    try:
        doc  = fitz.open(stream=file_bytes, filetype="pdf")
        page = doc[0]
        best, best_px = None, 0
        for img in page.get_images(full=True):
            try:
                bi  = doc.extract_image(img[0])
                w, h = bi["width"], bi["height"]
                if _is_portrait(w, h):
                    px = w * h
                    if px > best_px: best, best_px = bi["image"], px
            except Exception:
                continue
        if best: log.info("✓ Photo from PDF (%d px)", best_px)
        return best
    except Exception as e:
        log.error("Photo PDF error: %s", e)
        return None


def extract_photo_from_docx(file_bytes: bytes):
    try:
        best, best_px = None, 0
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            media = [f for f in z.namelist()
                     if f.startswith("word/media/") and
                     f.lower().rsplit(".", 1)[-1] in ("jpg", "jpeg", "png", "bmp")]
            for path in media:
                img_bytes = z.read(path)
                ext = "." + path.rsplit(".", 1)[-1].lower()
                w, h = _img_dimensions(img_bytes, ext)
                if _is_portrait(w, h):
                    px = w * h
                    if px > best_px: best, best_px = img_bytes, px
        if best: log.info("✓ Photo from DOCX (%d px)", best_px)
        return best
    except Exception as e:
        log.error("Photo DOCX error: %s", e)
        return None