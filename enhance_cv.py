# ─────────────────────────────────────────────────────────────────────────────
# enhance_cv.py — LLM rewrite, DOCX generation, and API endpoints
# Imported by main.py:  import enhance_cv
# All symbols from main.py are available via:  from main import ...
# ─────────────────────────────────────────────────────────────────────────────
import re
import io
import json
import os
import logging
from fastapi import APIRouter, UploadFile, File, Form, Path as FPath
from fastapi.responses import StreamingResponse, JSONResponse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from azure.cosmos import CosmosClient, PartitionKey, exceptions as cosmos_exceptions
from openai import AzureOpenAI
# ── Logger ────────────────────────────────────────────────────────────────────
log = logging.getLogger("cv_booster")

# ── Load .env BEFORE reading env vars ────────────────────────────────────────
from dotenv import load_dotenv
from pathlib import Path
load_dotenv(Path(__file__).parent / ".env")   # ← AJOUTER CES 3 LIGNES ICI

# ── Azure OpenAI (TES variables .env) ────────────────────────────────────────
AZURE_DEPLOYMENT = (
    os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME") or
    os.getenv("AZURE_OPENAI_DEPLOYMENT", "gpt-4o-mini")
)
client = AzureOpenAI(
    api_version    = os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview"),
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT", ""),
    api_key        = os.getenv("AZURE_OPENAI_API_KEY", ""),
)

# ── APIRouter — no circular import with main.py ───────────────────────────────
enhance_router = APIRouter(tags=["CV Booster"])

# ── Sentence Transformers — lazy import to avoid startup crash ────────────────
try:
    from sentence_transformers import SentenceTransformer, util as st_util
    _embedder   = SentenceTransformer("all-MiniLM-L6-v2")
    SEMANTIC_OK = True
except Exception:
    _embedder   = None
    SEMANTIC_OK = False

FITZ_OK = True  # pymupdf optional

# ── Logger ────────────────────────────────────────────────────────────────────
log = logging.getLogger("cv_booster")

# ── Azure OpenAI (TES variables .env) ────────────────────────────────────────
AZURE_DEPLOYMENT = (
    os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME") or
    os.getenv("AZURE_OPENAI_DEPLOYMENT", "gpt-4o-mini")
)
client = AzureOpenAI(
    api_version    = os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview"),
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT", ""),
    api_key        = os.getenv("AZURE_OPENAI_API_KEY", ""),
)

# ── CosmosDB (TES variables .env) ─────────────────────────────────────────────
COSMOS_ENDPOINT  = os.getenv("AZURE_COSMOS_ENDPOINT", "")
COSMOS_KEY       = os.getenv("AZURE_COSMOS_KEY", "")
COSMOS_DB        = os.getenv("AZURE_COSMOS_DATABASE_NAME", "EduTech_AI_Production")
COSMOS_CONTAINER = os.getenv("AZURE_COSMOS_USERS_CONTAINER", "users")

# SEMANTIC_OK defined above

# ── Helpers defined locally (same as save_cv.py) ─────────────────────────────
# _clean_noise, _has_real_content, parse_cv_sections, detect_domain
# are defined below in this file

# ── CV extraction ─────────────────────────────────────────────────────────────
from cv_extraction import (
    extract_text_from_pdf, extract_text_from_docx,
    extract_photo_from_pdf, extract_photo_from_docx,
)

# ── XAI explanation ───────────────────────────────────────────────────────────
from explain_ats import explain_ats_score

# ── Platform data ─────────────────────────────────────────────────────────────
from db_platform import get_platform_data_or_fallback, fetch_recommendations

# ── PLATFORM_DATA fallback (était dans le main.py de la collègue) ─────────────
PLATFORM_DATA = {
    "quiz": {"domain": "Data Analytics & BI", "score": 82, "level": "Advanced"},
    "labs": [],
    "certifications": [],
}

# ── ATS scoring (depuis le main.py de la collègue — recopié ici) ──────────────
_FALLBACK_KEYWORDS = [
    "machine learning", "data", "python", "analysis", "analytics",
    "artificial intelligence", "sql", "statistics", "model", "visualization",
    "engineering", "project", "development", "bachelor", "master",
    "scikit", "pandas", "tensorflow", "docker", "git", "power bi",
    "cloud", "aws", "pipeline", "dashboard",
]

_DOMAIN_KW_CACHE: dict = {}
# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 — HELPERS MANQUANTS (detect_domain, parse_cv_sections, etc.)
# ═════════════════════════════════════════════════════════════════════════════

# ── Domain detection ──────────────────────────────────────────────────────────
_DOMAIN_KEYWORDS = {
    "Data Science":       ["data science", "machine learning", "deep learning", "scikit", "tensorflow", "pytorch", "nlp", "neural network", "kaggle"],
    "Data Analytics & BI":["power bi", "tableau", "dashboard", "excel", "reporting", "kpi", "analytics", "looker", "business intelligence"],
    "Data Engineering":   ["spark", "airflow", "kafka", "etl", "pipeline", "dbt", "databricks", "bigquery", "redshift", "data warehouse"],
    "Cloud & DevOps":     ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ci/cd", "jenkins", "devops", "cloud"],
    "Software Engineering":["java", "spring", "c++", "c#", ".net", "microservices", "rest api", "backend", "frontend", "react", "angular"],
    "Cybersecurity":      ["cybersecurity", "security", "penetration", "siem", "firewall", "compliance", "iso 27001", "soc"],
    "AI & ML":            ["artificial intelligence", "llm", "gpt", "transformer", "rag", "generative", "prompt engineering", "langchain"],
    "NLP":                ["nlp", "natural language", "text mining", "sentiment", "named entity", "spacy", "bert", "hugging face"],
    "Web Development":    ["html", "css", "javascript", "typescript", "react", "vue", "node", "django", "flask", "fastapi"],
    "Finance":            ["finance", "accounting", "financial", "budget", "audit", "investment", "banking"],
    "Marketing":          ["marketing", "seo", "sem", "social media", "campaign", "digital marketing", "content"],
    "Product Management": ["product manager", "product owner", "roadmap", "agile", "scrum", "backlog", "user story"],
}

def detect_domain(cv_text: str) -> str:
    """Detect the professional domain from CV text."""
    text_lo = cv_text.lower()
    scores = {}
    for domain, keywords in _DOMAIN_KEYWORDS.items():
        scores[domain] = sum(1 for kw in keywords if kw in text_lo)
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "Software Engineering"


# ── CV section parser ─────────────────────────────────────────────────────────
def parse_cv_sections(cv_text: str) -> dict:
    """Parse raw CV text into sections dict for ATS scoring."""
    lines   = cv_text.split("\n")
    sections: dict = {}
    current = "header"
    sections[current] = []

    for line in lines:
        s = line.strip()
        if not s:
            continue
        matched = next(
            (k for k, p in SECTION_HEADERS_RE_ENH.items() if p.match(s)),
            None
        )
        if matched:
            current = matched
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, [])
            sections[current].append(s)

    return sections


# ── Text cleaning ─────────────────────────────────────────────────────────────
_NOISE_LINES = {
    "curriculum vitae", "cv", "resume", "résumé",
    "page 1", "page 2", "page 3",
    "confidential", "confidentiellement",
}

def _clean_noise(text: str) -> str:
    """Remove noise lines and normalize whitespace."""
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        s = line.strip()
        # Skip pure noise
        if s.lower() in _NOISE_LINES:
            continue
        # Skip lines that are just repeated symbols
        if s and len(set(s.replace(" ", ""))) <= 2 and len(s) > 3:
            continue
        cleaned.append(line)
    result = "\n".join(cleaned)
    # Collapse 3+ blank lines into 2
    import re as _re
    result = _re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


# ── Content validation ────────────────────────────────────────────────────────
_TEMPLATE_SIGNALS = [
    "your name here", "full name", "your email", "your phone",
    "job title here", "company name here", "degree title",
    "institution name", "your skills here", "add your",
    "click here to", "type your", "enter your",
    "[name]", "[email]", "[phone]", "[title]",
]

def _has_real_content(text: str) -> tuple[bool, str]:
    """
    Returns (True, "") if the CV has real content.
    Returns (False, reason) if it looks like an empty template.
    """
    text_lo = text.lower()

    # Check for template signals
    template_hits = sum(1 for sig in _TEMPLATE_SIGNALS if sig in text_lo)
    if template_hits >= 3:
        return False, (
            "Your CV appears to be an unfilled template "
            f"({template_hits} placeholder fields detected). "
            "Please fill in your real information before uploading."
        )

    # Must have some minimal length
    words = text.split()
    if len(words) < 30:
        return False, (
            f"Your CV is too short ({len(words)} words). "
            "Please upload a complete CV."
        )

    # Must contain at least one email-like or name-like token
    import re as _re
    has_email = bool(_re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text))
    has_name  = bool(_re.search(r"[A-Z][a-z]+ [A-Z][a-z]+", text))
    if not has_email and not has_name:
        return False, (
            "Could not find a name or email in your CV. "
            "Please make sure you uploaded the correct file."
        )

    return True, ""
SECTION_HEADERS_RE_ENH = {
    "profile":        __import__('re').compile(r"^(profile|profil|about|objective|summary|résumé|resume|présentation|professional summary|about me)$", __import__('re').I),
    "education":      __import__('re').compile(r"^(education|formation|éducation|études|diplômes?)$", __import__('re').I),
    "experience":     __import__('re').compile(r"^(experience|expérience|professional experience|work experience|internships?)$", __import__('re').I),
    "skills":         __import__('re').compile(r"^(skills?|technical skills?|compétences?|technologies|expertise)$", __import__('re').I),
    "projects":       __import__('re').compile(r"^(projects?|projets?|portfolio)$", __import__('re').I),
    "certifications": __import__('re').compile(r"^(certif\w*|badges?|awards?|training)$", __import__('re').I),
    "languages":      __import__('re').compile(r"^(languages?|langues?)$", __import__('re').I),
}

def get_domain_keywords(domain: str) -> list:
    if domain in _DOMAIN_KW_CACHE:
        return _DOMAIN_KW_CACHE[domain]
    try:
        resp = client.chat.completions.create(
            model=AZURE_DEPLOYMENT,
            messages=[{"role": "user", "content": (
                f"List the 60 most important ATS keywords for a '{domain}' professional CV.\n"
                f"Return ONLY a valid JSON array of lowercase strings.\n"
                f'Example: ["python", "sql", "data pipeline"]'
            )}],
            max_tokens=600, temperature=0,
        )
        raw = resp.choices[0].message.content.strip()
        raw = re.sub(r"```json|```", "", raw).strip()
        m   = re.search(r"\[.*\]", raw, re.DOTALL)
        keywords = json.loads(m.group()) if m else _FALLBACK_KEYWORDS
        keywords = [str(k).lower().strip() for k in keywords if k]
        _DOMAIN_KW_CACHE[domain] = keywords
        return keywords
    except Exception:
        return _FALLBACK_KEYWORDS

def calculate_ats_score(sections: dict, full_text: str, domain_keywords: list | None = None) -> dict:
    scores  = {}
    text_lo = full_text.lower()
    important = ["profile", "education", "experience", "skills", "projects"]
    found     = [s for s in important if sections.get(s)]
    scores["sections"] = {"score": round(len(found) / len(important) * 20), "max": 20, "detail": f"{len(found)}/{len(important)} key sections"}
    kw_list = domain_keywords if domain_keywords else _FALLBACK_KEYWORDS
    kw_found = [kw for kw in kw_list if kw in text_lo]
    kw_score = min(round(len(kw_found) / max(len(kw_list) * 0.6, 1) * 8), 8)
    scores["keywords"]  = {"score": kw_score, "max": 8, "detail": f"{len(kw_found)}/{len(kw_list)} keywords", "matched": [{"keyword": k, "similarity": 1.0} for k in kw_found]}
    exp_lines   = sections.get("experience", [])
    has_dates   = any(re.search(r"\d{4}", l) for l in exp_lines)
    has_bullets = any(l.strip().startswith(("-", "•")) for l in exp_lines)
    scores["experience"] = {"score": min((9 if exp_lines else 0) + (5 if has_dates else 0) + (4 if has_bullets else 0), 18), "max": 18, "detail": f"{len(exp_lines)} lines"}
    total_lines = sum(len(v) for v in sections.values())
    scores["length"] = {"score": 12 if 20 <= total_lines <= 70 else (4 if total_lines < 20 else 8), "max": 12, "detail": f"{total_lines} lines"}
    combined = text_lo[:500]
    has_phone = bool(re.search(r"\+?\d[\d\s\-]{7,}", combined))
    scores["contact"] = {"score": (3 if "@" in combined else 0) + (3 if has_phone else 0) + (2 if "linkedin" in combined else 0), "max": 8, "detail": f"Email:{'@' in combined}"}
    profile_text  = " ".join(sections.get("profile", []) + sections.get("summary", [])).strip()
    summary_score = (5 if profile_text else 0) + (3 if len(profile_text.split()) >= 20 else 0) + (2 if any(kw in profile_text.lower() for kw in ["experienced", "motivated", "skilled", "professional"]) else 0)
    scores["summary"] = {"score": min(summary_score, 10), "max": 10, "detail": f"{len(profile_text.split())} words"}
    skill_kws = ["python", "java", "sql", "javascript", "docker", "git", "tensorflow", "pandas", "aws", "machine learning"]
    skill_count = sum(1 for s in skill_kws if s in text_lo)
    scores["skills"] = {"score": min(round(skill_count / 5 * 18), 18), "max": 18, "detail": f"{skill_count} skills"}
    lang_lines = sections.get("languages", [])
    scores["languages"] = {"score": min((3 if lang_lines else 0) + (2 if any(m in " ".join(lang_lines).lower() for m in ["native", "fluent", "b2", "c1"]) else 0) + (1 if len(lang_lines) >= 2 else 0), 6), "max": 6, "detail": f"{len(lang_lines)} lang lines"}
    total = sum(v["score"] for v in scores.values())
    return {"total": total, "max": 100, "breakdown": scores}








def _get_cosmos_container():
    c = CosmosClient(COSMOS_ENDPOINT, credential=COSMOS_KEY)
    db = c.create_database_if_not_exists(COSMOS_DB)
    return db.create_container_if_not_exists(
        id=COSMOS_CONTAINER,
        partition_key=PartitionKey(path="/id"),
    )

def extract_cv_structured(llm_text: str) -> dict:
    """Appelle Azure OpenAI pour extraire les champs structurés depuis le CV amélioré."""
    prompt = f"""
You are a CV data extractor. From the CV text below, extract ONLY these fields as JSON:
{{
  "first_name": "string",
  "last_name": "string",
  "email": "string",
  "linkedin": "string (URL or username)",
  "role": "current/target job title",
  "seniority": "Junior | Mid | Senior | Lead | Executive",
  "years_experience": "number as string",
  "industry": "string",
  "education": "highest degree + institution",
  "skills": "comma-separated top skills",
  "summary": "2-3 sentence professional summary",
  "bullets": "top 3 achievement bullets, semicolon-separated"
}}
Return ONLY valid JSON. No markdown, no explanation.

CV TEXT:
{llm_text[:4000]}
"""
    resp = client.chat.completions.create(
        model=AZURE_DEPLOYMENT,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=600,
        temperature=0,
    )
    raw = resp.choices[0].message.content.strip()
    # Nettoyer les backticks si présents
    raw = raw.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)


def save_cv_to_cosmos(user_id: str, cv_structured: dict):
    """Upsert le profil utilisateur dans CosmosDB."""
    try:
        container = _get_cosmos_container()
        # Lire l'existant pour ne pas écraser des champs déjà remplis
        try:
            existing = container.read_item(item=user_id, partition_key=user_id)
        except cosmos_exceptions.CosmosResourceNotFoundError:
            existing = {}

        doc = {
            "id":         user_id,
            "first_name": cv_structured.get("first_name") or existing.get("first_name"),
            "last_name":  cv_structured.get("last_name")  or existing.get("last_name"),
            "email":      cv_structured.get("email")      or existing.get("email"),
            "linkedin":   cv_structured.get("linkedin")   or existing.get("linkedin"),
            "role":       cv_structured.get("domain", ""),
            "seniority":  cv_structured.get("level", ""),
            "years_exp":  cv_structured.get("years_experience", ""),
            "industry":   cv_structured.get("industry", ""),
            "education":  cv_structured.get("education", ""),
            "skills":     cv_structured.get("skills", ""),
            "summary":    cv_structured.get("summary", ""),
            "bullets":    cv_structured.get("bullets", ""),   
            "ats_score_before":  cv_structured.get("ats_score_before", 0),   
            "ats_score_after":  cv_structured.get("ats_score_after", 0),   
        }
        container.upsert_item(doc)
        log.info("✅ CosmosDB upsert — user_id: %s", user_id)
    except Exception as e:
        log.error("❌ CosmosDB save failed: %s", str(e))
# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5 — LLM CV GENERATION
# ═════════════════════════════════════════════════════════════════════════════

def build_platform_summary(platform: dict, include_quiz=True,
                           include_lab_ids=None, include_cert_ids=None) -> str:
    lines = []
    if include_quiz and platform.get("quiz"):
        q = platform["quiz"]
        lines += [
            "PLATFORM POSITIONING QUIZ:",
            f"  - Domain: {q['domain']}",
            f"  - Score: {q['score']}% — Level: {q['level']}",
            "  - Add as validated skill level indicator in the CV profile section", "",
        ]
    labs = [l for l in platform.get("labs", [])
            if include_lab_ids is None or l["id"] in include_lab_ids]
    if labs:
        lines.append("PLATFORM LABS — ADD EACH ONE AS A PROJECT ENTRY UNDER THE PROJECTS SECTION (not certifications):")
        for l in labs:
            lines.append(f"  - {l['title']} | Completed: {l['date']} | Score: {l['score']}/100")
        lines.append("")
    certs = [c for c in platform.get("certifications", [])
             if include_cert_ids is None or c["id"] in include_cert_ids]
    if certs:
        lines.append("VALIDATED CERTIFICATIONS (issued by official organizations):")
        for c in certs:
            lines.append(f"  - {c['title']} | Issued by: {c['org']} | Date: {c['date']}")
        lines.append("")
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# Skill inference from platform certs + labs (deterministic — no LLM needed)
# ─────────────────────────────────────────────────────────────────────────────

# ── Title keyword → skill (exact, single mapping) ────────────────────────────
# Each keyword found in a lab/cert title adds ONLY the skills directly named
# by that keyword. No speculation, no extras beyond what the title states.

_TITLE_KEYWORD_SKILLS: dict[str, list[str]] = {
    "machine learning":  ["Machine Learning", "Scikit-learn"],
    "python":            ["Python"],
    "sql":               ["SQL"],
    "advanced sql":      ["SQL", "Window Functions", "CTEs", "Query Optimization"],
    "power bi":          ["Power BI", "DAX"],
    "tableau":           ["Tableau"],
    "aws":               ["AWS", "Cloud Computing"],
    "azure":             ["Microsoft Azure", "Cloud Computing"],
    "google cloud":      ["Google Cloud"],
    "data analytics":    ["Data Analytics", "SQL"],
    "data analysis":     ["Data Analysis", "Python"],
    "data science":      ["Data Science", "Python"],
    "data engineering":  ["Data Engineering"],
    "deep learning":     ["Deep Learning", "Neural Networks"],
    "nlp":               ["NLP"],
    "tensorflow":        ["TensorFlow"],
    "pytorch":           ["PyTorch"],
    "spark":             ["Apache Spark", "PySpark"],
    "docker":            ["Docker"],
    "git":               ["Git"],
    "excel":             ["Excel"],
    "tableau":           ["Tableau"],
    "visualization":     ["Data Visualization"],
    "dashboard":         ["Dashboard Design"],
    "statistics":        ["Statistics"],
    "r programming":     ["R"],
    "cybersecurity":     ["Cybersecurity"],
    "cloud":             ["Cloud Computing"],
    "pandas":            ["Pandas"],
    "numpy":             ["NumPy"],
    "matplotlib":        ["Matplotlib"],
    "jupyter":           ["Jupyter Notebook"],
}


def infer_skills_from_platform(
    platform: dict,
    include_lab_ids: list | None,
    include_cert_ids: list | None,
) -> list[str]:
    """
    Extract skills directly from lab/cert titles — only what is explicitly
    named in the title. No speculation beyond the title keywords.
    Keywords are matched longest-first to avoid 'sql' matching 'advanced sql'.
    """
    skills: set[str] = set()

    # Sort keywords longest-first so "advanced sql" matches before "sql"
    sorted_keywords = sorted(_TITLE_KEYWORD_SKILLS.keys(), key=len, reverse=True)

    def extract_from_title(title: str):
        title_lo = title.lower()
        matched_ranges: list[tuple[int, int]] = []
        for kw in sorted_keywords:
            pos = title_lo.find(kw)
            if pos == -1:
                continue
            end = pos + len(kw)
            # Skip if this range is already covered by a longer match
            if any(s <= pos and end <= e for s, e in matched_ranges):
                continue
            matched_ranges.append((pos, end))
            skills.update(_TITLE_KEYWORD_SKILLS[kw])

    certs = [c for c in platform.get("certifications", [])
             if include_cert_ids is None or c["id"] in include_cert_ids]
    for cert in certs:
        extract_from_title(cert["title"])
        log.info("  cert '%s' → %s", cert["title"],
                 [s for kw in sorted_keywords
                  if kw in cert["title"].lower()
                  for s in _TITLE_KEYWORD_SKILLS[kw]])

    labs = [l for l in platform.get("labs", [])
            if include_lab_ids is None or l["id"] in include_lab_ids]
    for lab in labs:
        extract_from_title(lab["title"])
        log.info("  lab  '%s' → %s", lab["title"],
                 [s for kw in sorted_keywords
                  if kw in lab["title"].lower()
                  for s in _TITLE_KEYWORD_SKILLS[kw]])

    result = sorted(skills)
    log.info("✓ Inferred %d skills from titles (certs:%d labs:%d)",
             len(result), len(certs), len(labs))
    return result



# ── Cert garble pre-processor ────────────────────────────────────────────────
_CERT_GARBLE_SIGNALS = [
    r"engineer.*engineer", r"learning.*learning", r"fundamentals.*fundamentals",
    r"\d{4}.*\d{4}.*\d{4}",
]
def _cert_block_is_garbled(cert_lines):
    blob = " ".join(cert_lines).lower()
    for p in _CERT_GARBLE_SIGNALS:
        if re.search(p, blob): return True
    for line in cert_lines:
        if len(re.findall(r"\b20\d{2}\b", line)) >= 2: return True
    return False
def _extract_cert_lines(cv_text):
    lines, cert_lines, in_cert = cv_text.split("\n"), [], False
    cert_re = re.compile(r"^(certif\w*|training|badges?|awards?)", re.I)
    end_re  = re.compile(r"^(langues?|languages?|compétences?|skills?|éducation|education"
                         r"|expérience|experience|projets?|projects?|profil|profile|summary)$", re.I)
    for line in lines:
        s = line.strip()
        if cert_re.match(s): in_cert = True; continue
        if in_cert:
            if end_re.match(s): break
            if s: cert_lines.append(s)
    return cert_lines
def fix_garbled_certifications(cv_text):
    cert_lines = _extract_cert_lines(cv_text)
    if not cert_lines or not _cert_block_is_garbled(cert_lines): return cv_text
    log.info("⚠  Cert block garbled — reconstructing via LLM")
    known_orgs = ["DataCamp","Amazon Web Services","AWS","NVIDIA","IBM","Google",
                  "Microsoft","Coursera","Udemy","LinkedIn Learning","Oracle","CompTIA"]
    raw_blob = "\n".join(cert_lines)
    prompt = (
        "The text is a certifications section from a 3-column PDF (cert name|org|year scrambled).\n"
        "Reconstruct every certification correctly.\n"
        "RULES: cert name can span multiple lines; known orgs: " + ", ".join(known_orgs) + ";\n"
        "never merge two cert names; never invent certs.\n"
        "Output ONLY bullet lines: - Certification Name · Org — Year\n\n"
        f"RAW:\n{raw_blob}"
    )
    try:
        resp = client.chat.completions.create(
            model=AZURE_DEPLOYMENT,
            messages=[{"role":"user","content":prompt}],
            temperature=0, max_tokens=600,
        )
        clean = resp.choices[0].message.content.strip()
        log.info("✓ Cert reconstruction done")
        lines, out, in_cert, replaced = cv_text.split("\n"), [], False, False
        cert_re2 = re.compile(r"^(certif\w*|training|badges?|awards?)", re.I)
        end_re2  = re.compile(r"^(langues?|languages?|compétences?|skills?|éducation|education"
                              r"|expérience|experience|projets?|projects?|profil|profile|summary)$", re.I)
        for line in lines:
            s = line.strip()
            if cert_re2.match(s) and not replaced:
                in_cert = True; out.append(line); out.append(clean); continue
            if in_cert:
                if end_re2.match(s): in_cert = False; replaced = True; out.append(line)
            else: out.append(line)
        return "\n".join(out)
    except Exception as e:
        log.error("Cert reconstruction failed: %s", e); return cv_text

def generate_cv_with_llm(
    cv_text: str,
    platform: dict,
    include_quiz=True,
    include_lab_ids=None,
    include_cert_ids=None,
    skipped_sections: list | None = None,
) -> str:
    skipped_sections = skipped_sections or []

    platform_summary = build_platform_summary(
        platform, include_quiz, include_lab_ids, include_cert_ids
    )

    # ── Inferred skills from platform certs + labs ────────────────────────────
    inferred_skills = infer_skills_from_platform(platform, include_lab_ids, include_cert_ids)
    skills_block = ""
    if inferred_skills and "skills" not in skipped_sections:
        skills_block = (
            "\n\nSKILLS INFERRED FROM PLATFORM CERTIFICATIONS & LABS "
            "(verified — include ALL of these in the SKILLS section):\n"
            + ", ".join(inferred_skills)
        )

    # ── Hard OMIT block — injected at end of user message for max LLM attention
    _LABEL_MAP = {
        "languages":      "LANGUAGES",
        "education":      "EDUCATION",
        "experience":     "PROFESSIONAL EXPERIENCE",
        "projects":       "PROJECTS",
        "skills":         "SKILLS",
        "certifications": "CERTIFICATIONS & TRAINING",
        "profile":        "PROFILE",
    }
    omit_block = ""
    if skipped_sections:
        omit_labels = [_LABEL_MAP[s] for s in skipped_sections if s in _LABEL_MAP]
        if omit_labels:
            omit_block = (
                "\n\n════════════════════════════════\n"
                "SECTIONS THE USER CHOSE TO SKIP — OMIT COMPLETELY\n"
                "════════════════════════════════\n"
                "Do NOT write these sections. Do NOT write their header. "
                "Do NOT write any placeholder or 'Omitted' text. "
                "Simply exclude them as if they do not exist:\n"
                + "\n".join(f"  ✗ {lbl}" for lbl in omit_labels)
                + "\n════════════════════════════════"
            )
            log.info("🚫 LLM instructed to omit: %s", omit_labels)

    # ── Detect CV language ───────────────────────────────────────────────────
    _fr_words = ["et","ou","les","des","dans","pour","avec","sur","une","est",
                 "par","je","mon","ma","du","au","expérience","formation","compétences","langues"]
    _cv_lower_lang = cv_text.lower()
    _fr_count = sum(1 for w in _fr_words if f" {w} " in _cv_lower_lang)
    _cv_language = "French" if _fr_count >= 4 else "English"
    _language_instruction = (
        f"CRITICAL: The entire CV MUST be written in {_cv_language} — "
        f"the same language as the original CV. Every section header, bullet, "
        f"description, and skill category label must be in {_cv_language}. "
        f"Do NOT switch languages mid-document."
    )
    log.info("✓ CV language detected: %s (fr_signals=%d)", _cv_language, _fr_count)

    # ── Extract quiz domain for profile instruction ───────────────────────────
    quiz_domain = (platform.get("quiz") or {}).get("domain", "")
    quiz_level  = (platform.get("quiz") or {}).get("level", "")
    domain_instruction = ""
    if quiz_domain:
        domain_instruction = (
            f"- The candidate's validated domain is **{quiz_domain}** "
            f"(level: {quiz_level}) — write the profile around THIS domain\n"
            f"- Use keywords relevant to {quiz_domain}, not generic unrelated keywords\n"
        )

    system_prompt = f"""You are an expert CV writer specializing in ATS-optimized resumes for tech and data science roles.

{_language_instruction}

Your task: generate a complete, professional, ATS-friendly CV combining:
1. The candidate's original CV — use ALL real information
2. Their platform learning data (certifications, lab projects, assessed skills)

OUTPUT FORMAT (strictly follow):
- First line = candidate full name only
- Then EXACT section headers in ALL CAPS on their own line
- Blank line between sections
- No markdown, no code blocks, no explanations — CV text only
- Use **text** markers ONLY for: job titles, company names, degree titles, project names, skill category labels
- NEVER wrap dates, years, or numbers in ** markers

CONTACT SECTION:
- All on ONE line separated by  |  symbol

PROFILE SECTION:
- 3-4 strong sentences: who you are, what you do, what you bring
{domain_instruction}- Integrate platform strengths naturally
- PRESERVE all quantified facts: academic rank ("major de promotion" / "top of class"),
  number of experiences, specific metrics, internship target window, dual-degree details.
  These are ATS keywords — never drop or paraphrase them away.

EDUCATION SECTION:
- Format: **Degree Title**\t2020 - 2023
- Next line: Institution name, City
- ONLY include if education data exists in the CV — OMIT entirely if not found

PROFESSIONAL EXPERIENCE SECTION:
Format: **Job Title | Company Name**\tDate

Order all experiences by date in descending order (most recent first, oldest last)

 Bullet points starting with "- " using strong action verbs

ONLY include this section if experience data exists in the CV — OMIT entirely if not found

PROJECTS SECTION:
- ONLY include projects that EXIST in the original CV or in the platform labs provided
- NEVER invent, imagine, or create projects that are not explicitly in the source data
- Include ALL projects from the CV — do not truncate or omit any
- PLATFORM LABS → PROJECTS: Every platform lab MUST appear as a project entry here.
  Labs are NEVER certifications. NEVER put a lab under CERTIFICATIONS.
- If no projects exist in the CV AND no platform labs were provided → OMIT this section entirely
- Merge CV projects + platform labs into one unified PROJECTS section
- Sort by date descending (most recent first)
- Remove duplicates
- CRITICAL — ONE PROJECT PER ENTRY: Every project and lab is its own separate paragraph.
  NEVER merge two project names into one line. Merging is a HARD ERROR.
- Format each project/lab as:
  **Project Name** — Year
  One concise line: what was built and which technologies were used.
- ONLY include this section if at least one real project or lab exists in the data provided

SKILLS SECTION:
- Combine skills from: (a) original CV, (b) platform inferred skills provided below
- CRITICAL: include EVERY skill from the original CV — do not drop any.
  ETL, RAG, diffusion models, prompt engineering, OCR, agents — ALL must appear.
- Group into REAL named categories:
    **Programming Languages:** Python, Java, SQL, JavaScript, ...
    **Data & Machine Learning:** Scikit-learn, TensorFlow, PyTorch, Pandas, ETL, LLM, ...
    **NLP & Generative AI:** NLP, LLM, RAG, Diffusion models, Prompt engineering, Agents, ...
    **NLP & OCR:** OCR, Text extraction, Structured extraction, ...
    **Databases & SQL:** SQL, PostgreSQL, BigQuery, ...
    **Cloud & DevOps:** AWS, Azure, Docker, CI/CD, Jenkins, Airflow, Git, ...
    **Data Visualization:** Power BI, Tableau, Matplotlib, Excel/VBA, ...
    **Tools & Productivity:** Git, Jupyter Notebook, Linux, ...
- Use the EXACT category name that fits — NEVER write the word "Category" as a label
- Only include a category if it has at least one real skill to list
- Never discard any skill — if it doesn't fit, add it to the closest category

CERTIFICATIONS & TRAINING SECTION:
You are a professional CV writer and technical recruiter.

Your task is to improve and structure the CERTIFICATIONS section of a CV.

Instructions:

1. Collect all certifications mentioned in the CV, including:
   - official certifications
   - platform-based certificates (Coursera, Udemy, AWS, LinkedIn Learning, etc.)

2. Merge all certifications into **one unified CERTIFICATIONS section** — do not separate platform certificates from official ones.

3. Remove duplicates if the same certification appears more than once.

4. **Sort all certifications by date in descending order** (most recent first, oldest last), regardless of whether they come from the CV or platforms.

5. Format each certification as follows:

CERTIFICATIONS

- **Certification Name** · Issued by Organization — Date

6. Only include this section if **at least one certification exists**.

Example:

CERTIFICATIONS

- **Machine Learning Specialization** · Coursera — 2025  
- **AWS Solutions Architect Associate** · Amazon Web Services — 2024  
- **Python for Data Science** · Udemy — 2023  
- **Power BI Data Analyst** · Microsoft — 2023

Important:  
All certifications, whether from the CV or platforms, must be **ordered from most recent to oldest**.


LANGUAGES SECTION:
- Include if languages are present anywhere in the provided text (CV or supplementary data)
- Format: Language: Level | Language: Level
- If none found → OMIT this section completely

ABSOLUTE RULES:
- NEVER write placeholder text like 'Degree Title', 'Institution name', 'City', 'Job Title Here'
- NEVER write the word 'Omitted' or 'N/A' or any placeholder in any section
- NEVER invent information not in the source CV or platform data
- NEVER invent projects, experiences, or achievements — only use what is explicitly provided
- NEVER use the word 'Category' as a skill category label
- NEVER Write only one word for job title 
- If a section has no real data → OMIT it entirely (no header, no content)
- If the CV is clearly an empty template → respond only with: ERROR_TEMPLATE
- Aim for 50-65 lines total
- Use \\t (tab character) to push dates to the right
"""

    import re as _re
    _pf = []
    _cvlo = cv_text.lower()
    if "major de promotion" in _cvlo: _pf.append("major de promotion (top of class)")
    if "top of class" in _cvlo: _pf.append("top of class")
    _em = _re.search(r"(\d+)\+?\s*(expériences?|experiences?|stages?|internships?)", _cvlo)
    if _em: _pf.append(f"{_em.group(1)}+ significant professional experiences")
    _dm = _re.search(r"(f[eé]vrier|february|mars|march|avril|april|mai|may|juin|june|juillet|july).*?(\d{{4}})", _cvlo)
    if _dm:
        s, e = max(0,_dm.start()-5), min(len(_cvlo),_dm.end()+20)
        _pf.append(f"seeking internship: {cv_text[s:e].strip()}")
    _facts_block = ("\n\nPROFILE FACTS — MUST ALL APPEAR IN THE PROFILE SECTION:\n"
                    + "\n".join(f"  • {f}" for f in _pf)) if _pf else ""

    user_prompt = f"""CANDIDATE'S ORIGINAL CV:
---
{cv_text}
---

PLATFORM LEARNING DATA TO INTEGRATE:
---
{platform_summary}
---
{skills_block}
{_facts_block}
{omit_block}

Generate the complete ATS-optimized CV now."""

    log.info("📄 Sending CV to GPT for rewrite (%d chars, skipping: %s)...",
             len(cv_text), skipped_sections)
    resp = client.chat.completions.create(
        model=AZURE_DEPLOYMENT,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
        temperature=0.3,
        max_tokens=3000,
    )
    result = resp.choices[0].message.content.strip()
    log.info("✓ GPT rewrite complete (%d chars)", len(result))
    return result


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PARSE LLM OUTPUT
# ═════════════════════════════════════════════════════════════════════════════

LLM_SECTION_HEADERS = {
    "contact":        re.compile(r"^CONTACT$", re.I),
    "profile":        re.compile(
        r"^(PROFILE|SUMMARY|PROFESSIONAL SUMMARY|ABOUT|ABOUT ME"
        r"|PROFIL|RÉSUMÉ|RESUME|PRÉSENTATION|À PROPOS)$", re.I),
    "education":      re.compile(
        r"^(EDUCATION|ÉDUCATION|FORMATION|ÉTUDES|PARCOURS ACADÉMIQUE"
        r"|PARCOURS SCOLAIRE|DIPLÔMES?)$", re.I),
    "experience":     re.compile(
        r"^(PROFESSIONAL EXPERIENCE|EXPERIENCE|WORK EXPERIENCE|WORK HISTORY"
        r"|EXPÉRIENCE PROFESSIONNELLE|EXPÉRIENCE|EXPÉRIENCES PROFESSIONNELLES"
        r"|PARCOURS PROFESSIONNEL|STAGES?)$", re.I),
    "projects":       re.compile(
        r"^(PROJECTS|ACADEMIC PROJECTS|PROJETS?|PROJETS ACADÉMIQUES"
        r"|RÉALISATIONS?)$", re.I),
    "skills":         re.compile(
        r"^(SKILLS|TECHNICAL SKILLS|COMPÉTENCES?|COMPÉTENCES TECHNIQUES"
        r"|TECHNOLOGIES|EXPERTISE)$", re.I),
    "certifications": re.compile(
        r"^(CERTIFICATIONS? & TRAINING|CERTIFICATIONS?|TRAINING"
        r"|CERTIFICATIONS? AND TRAINING|CERTIFICATIONS? & FORMATIONS?"
        r"|FORMATIONS? COMPLÉMENTAIRES?)$", re.I),
    "languages":      re.compile(
        r"^(LANGUAGES?|LANGUES?|COMPÉTENCES? LINGUISTIQUES?)$", re.I),
}


# Lines the LLM sometimes outputs when it has no real data — strip them all
_PLACEHOLDER_LINES = {
    "degree title", "institution name", "city", "institution name, city",
    "omitted", "n/a", "not applicable", "job title", "company name",
    "job title here", "company name here", "your job title", "language: level",
    "language : level", "none", "—", "-", "category:", "category",
}

def _is_placeholder_line(line: str) -> bool:
    """Return True if a line is a known placeholder that should be stripped."""
    stripped = line.strip().lower()
    # Exact match
    if stripped in _PLACEHOLDER_LINES:
        return True
    # Looks like template: "Degree Title\t2020 - 2023"
    if re.match(r"^(degree title|institution name|job title|company name)", stripped, re.I):
        return True
    return False


def sanitize_parsed_sections(sections: dict, skipped: list) -> dict:
    """
    Two-pass cleanup after LLM output is parsed:
    1. Remove sections the user explicitly skipped (hard delete)
    2. Strip placeholder lines from every remaining section
    3. Delete any section that ends up empty after stripping
    """
    cleaned = {}
    for key, lines in sections.items():
        # Hard-delete skipped sections regardless of content
        if key in skipped:
            log.info("🗑  Removing skipped section from parsed output: %s", key)
            continue
        # Strip placeholder lines
        real_lines = [l for l in lines if not _is_placeholder_line(l)]
        # Keep section only if it has real content
        if real_lines:
            cleaned[key] = real_lines
        else:
            log.info("🗑  Dropping empty section after placeholder strip: %s", key)
    return cleaned


def parse_llm_output(text: str, skipped_sections: list | None = None) -> dict:
    skipped_sections = skipped_sections or []
    lines   = [l.strip() for l in text.split("\n")]
    result  = {"name": "", "sections": {"contact": []}}
    current = "contact"
    first   = True
    for line in lines:
        if not line: continue
        if first: result["name"] = line; first = False; continue
        clean   = line.rstrip(":").strip()
        matched = next((k for k, p in LLM_SECTION_HEADERS.items() if p.match(clean)), None)
        if matched:
            current = matched
            result["sections"].setdefault(current, [])
        else:
            v = line.rstrip()
            if v:
                result["sections"].setdefault(current, [])
                result["sections"][current].append(v)
    if not result["sections"].get("contact"):
        result["sections"].pop("contact", None)
    # Sanitize: remove skipped sections + strip placeholder lines
    result["sections"] = sanitize_parsed_sections(result["sections"], skipped_sections)
    return result


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 7 — JOB TITLE INFERENCE
# ═════════════════════════════════════════════════════════════════════════════

DOMAIN_TO_ROLE = {
    "data analytics & bi": "Data Analyst",   "data analytics": "Data Analyst",
    "business intelligence": "BI Analyst",   
    "data science": "Data Scientist",         "data engineering": "Data Engineer",
    "ai & ml": "AI Engineer",                "ai": "AI Engineer",
    "artificial intelligence": "AI Engineer","nlp": "NLP Engineer",
    "cloud computing": "Cloud Engineer",     "cloud": "Cloud Engineer",
    "machine learning": "Machine Learning Engineer","Java":"Software Engineer",
    "software engineering": "Software Engineer", "web development": "Web Developer",
    "cybersecurity": "Cybersecurity Specialist","cyber": "Cybersecurity Specialist",
    "devops": "DevOps Engineer",             "product": "Product Analyst",
    "finance": "Financial Analyst",          "marketing": "Marketing Analyst",
}
LEVEL_TO_SENIORITY = {
    "beginner":     "Junior",
    "intermediate": "Mid-Level",
    "advanced":     "Senior",
    "expert":       "Lead",
}


def infer_job_title(platform: dict, parsed_sections: dict, cv_domain: str = "", cv_stated_title: str = "") -> str:

    # Skills that should NEVER be treated as job titles
    SKILL_WORDS = {
        "python","java","sql","c++","c#","javascript","react","node","html","css",
        "aws","azure","gcp","docker","kubernetes","git","linux",
        "tensorflow","pytorch","pandas","numpy","matplotlib",
        "power bi","tableau","excel","spark","airflow"
    }

    # Role detection signals
    ROLE_SIGNALS = {

        "Machine Learning Engineer":[
            "machine learning","tensorflow","pytorch","scikit-learn","deep learning"
        ],

        "Data Scientist":[
            "data science","statistics","predictive modeling","pandas","numpy"
        ],

        "Data Engineer":[
            "spark","airflow","etl","data pipeline","dbt"
        ],

        "Data Analyst":[
            "power bi","tableau","dashboard","data analytics","excel","sql"
        ],

        "Software Engineer":[
            "java","c++","c#",".net","software engineer","software developer"
        ],

        "Backend Developer":[
            "django","flask","fastapi","node","spring boot","backend"
        ],

        "Frontend Developer":[
            "react","angular","vue","javascript","html","css"
        ],

        "Cloud Engineer":[
            "aws","azure","gcp","cloud computing"
        ],

        "DevOps Engineer":[
            "docker","kubernetes","ci/cd","terraform"
        ],

        "NLP Engineer":[
            "nlp","natural language processing","text mining"
        ]
    }

    # Combine all CV text
    cv_text = ""

    for section in parsed_sections.values():
        if isinstance(section, str):
            cv_text += section + " "
        elif isinstance(section, list):
            cv_text += " ".join(section) + " "

    cv_text = cv_text.lower()

    # ---------------------------
    # 1️⃣ Check stated title
    # ---------------------------
    stated = cv_stated_title.strip().lower()

    if stated and stated not in SKILL_WORDS and len(stated.split()) > 1:
        return stated.title()

    # ---------------------------
    # 2️⃣ Detect role from skills
    # ---------------------------
    role_scores = {}

    for role in ROLE_SIGNALS:
        role_scores[role] = 0

    for role, keywords in ROLE_SIGNALS.items():
        for k in keywords:
            if k in cv_text:
                role_scores[role] += 1

    best_role = max(role_scores, key=role_scores.get)

    if role_scores[best_role] > 0:
        return best_role

    # ---------------------------
    # 3️⃣ Fallback using domain
    # ---------------------------
    DOMAIN_TO_ROLE = {
        "java": "Software Engineer",
        "python": "Software Engineer",
        "machine learning": "Machine Learning Engineer",
        "data science": "Data Scientist",
        "data analytics": "Data Analyst",
        "cloud": "Cloud Engineer",
        "devops": "DevOps Engineer"
    }

    domain = cv_domain.lower()

    if domain in DOMAIN_TO_ROLE:
        return DOMAIN_TO_ROLE[domain]

    # ---------------------------
    # 4️⃣ Final fallback
    # ---------------------------
    return "Software Engineer"

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 8 — DOCX GENERATION HELPERS
# ═════════════════════════════════════════════════════════════════════════════

FONT  = "Times New Roman"
BLACK = RGBColor(0x00, 0x00, 0x00)
DGRAY = RGBColor(0x22, 0x22, 0x22)


def _add_bottom_border(para, color="000000", sz="8"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), sz)
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)

add_bottom_border = _add_bottom_border


def _set_right_tab(para, inches=7.1):
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(inches * 1440)))
    tabs.append(tab); pPr.append(tabs)

set_tab_stop_right = _set_right_tab


def _add_runs(p, text, size, color, font):
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if not part: continue
        bold = part.startswith("**") and part.endswith("**")
        r = p.add_run(part[2:-2] if bold else part)
        r.bold = bold; r.font.size = size; r.font.color.rgb = color; r.font.name = font


def _shading(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex); pPr.append(shd)


def _contact_str(parsed) -> str:
    cl = parsed["sections"].get("contact", [])
    if not cl: return ""
    return cl[0] if len(cl) == 1 else "  |  ".join(l.strip() for l in cl if l.strip())


def _render_line(doc, line, font, size, color, bullet_color=None, tab_inches=7.0):
    s  = line.strip()
    if not s: return
    bc = bullet_color or color
    ib = s.startswith("- ") or s.startswith("• ")
    clean = s[2:].strip() if ib else s
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    if ib:
        p.paragraph_format.left_indent       = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        rb = p.add_run("•  "); rb.font.size = size; rb.font.color.rgb = bc; rb.font.name = font
    if "\t" in clean:
        _set_right_tab(p, tab_inches)
        left, right = clean.split("\t", 1)
        _add_runs(p, left, size, color, font)
        p.add_run("\t").font.name = font
        _add_runs(p, right, size, color, font)
    else:
        _add_runs(p, clean, size, color, font)


def _body_line(doc, text):
    s  = text.strip()
    ib = s.startswith("- ") or s.startswith("• ")
    p  = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if ib:
        clean = s[2:].strip()
        p.paragraph_format.left_indent       = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        r = p.add_run("•  ")
        r.font.size = Pt(10.5); r.font.color.rgb = BLACK; r.font.name = FONT
        _add_runs(p, clean, Pt(10.5), BLACK, FONT)
    elif "\t" in s:
        _set_right_tab(p)
        left, right = s.split("\t", 1)
        _add_runs(p, left.strip(),  Pt(10.5), BLACK, FONT)
        p.add_run("\t").font.name = FONT
        _add_runs(p, right.strip(), Pt(10.5), BLACK, FONT)
    else:
        _add_runs(p, s, Pt(10.5), BLACK, FONT)
    return p

body_line      = _body_line
add_bold_parts = lambda p, text, size=Pt(10.5), color=None: _add_runs(p, text, size, color or BLACK, FONT)


def _section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    _add_bottom_border(p)
    r = p.add_run(title.upper())
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = BLACK; r.font.name = FONT
    return p

section_heading = _section_heading


SECTION_ORDER  = ["profile","education","experience","projects","skills","certifications","languages"]
SECTION_LABELS_EN = {
    "profile":        "Profile",
    "education":      "Education",
    "experience":     "Professional Experience",
    "projects":       "Projects",
    "skills":         "Skills",
    "certifications": "Certifications & Training",
    "languages":      "Languages",
}
SECTION_LABELS_FR = {
    "profile":        "Profil",
    "education":      "Éducation",
    "experience":     "Expérience Professionnelle",
    "projects":       "Projets",
    "skills":         "Compétences",
    "certifications": "Certifications & Formations",
    "languages":      "Langues",
}
# Default to English; overridden per-request when French CV detected (see generate_docx call)
SECTION_LABELS = SECTION_LABELS_EN


# ═════════════════════════════════════════════════════════════════════════════
# FORMAT 1 — ATS CLASSIC
# ═════════════════════════════════════════════════════════════════════════════

def _make_no_border_table(doc, col_widths_twips: list):
    ncols = len(col_widths_twips); total = sum(col_widths_twips)
    tbl_obj = doc.add_table(rows=1, cols=ncols); tbl_obj.style = "Table Grid"; tbl = tbl_obj._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(total)); tblW.set(qn("w:type"), "dxa"); tblPr.append(tblW)
    borders = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}"); b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0"); b.set(qn("w:space"), "0"); b.set(qn("w:color"), "auto"); borders.append(b)
    tblPr.append(borders)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths_twips: gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    tbl.insert(0, grid)
    cells = tbl_obj.rows[0].cells
    for cell, w in zip(cells, col_widths_twips):
        tcp = cell._tc.get_or_add_tcPr(); cw = OxmlElement("w:tcW"); cw.set(qn("w:w"), str(w)); cw.set(qn("w:type"), "dxa"); tcp.append(cw)
        bd = OxmlElement("w:tcBorders")
        for side in ["top","left","bottom","right"]:
            b = OxmlElement(f"w:{side}"); b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0"); b.set(qn("w:space"), "0"); b.set(qn("w:color"), "auto"); bd.append(b)
        tcp.append(bd)
    return tbl_obj


def generate_docx_from_llm(parsed: dict, photo_bytes=None, language: str = "English") -> bytes:
    labels = SECTION_LABELS_FR if language == "French" else SECTION_LABELS_EN
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.8)
        sec.left_margin = sec.right_margin = Inches(0.6)
    doc.styles["Normal"].font.name      = FONT
    doc.styles["Normal"].font.size      = Pt(10.5)
    doc.styles["Normal"].font.color.rgb = BLACK

    if photo_bytes:
        tbl = _make_no_border_table(doc, [7800, 1560])
        lc, rc = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
        pn = lc.paragraphs[0]; pn.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = pn.add_run(parsed["name"].upper())
        r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLACK; r.font.name = FONT
        jt = parsed.get("job_title", "")
        if jt:
            pj = lc.add_paragraph(); rj = pj.add_run(jt)
            rj.font.size = Pt(11); rj.font.color.rgb = DGRAY; rj.font.name = FONT; rj.italic = True
        pp = rc.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            pp.add_run().add_picture(io.BytesIO(photo_bytes), width=Inches(1.05))
        except Exception as e:
            log.error("Photo insert error: %s", e)
    else:
        pn = doc.add_paragraph(); pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pn.paragraph_format.space_before = Pt(0); pn.paragraph_format.space_after = Pt(2)
        r = pn.add_run(parsed["name"].upper())
        r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLACK; r.font.name = FONT
        jt = parsed.get("job_title", "")
        if jt:
            pj = doc.add_paragraph(); pj.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pj.paragraph_format.space_before = Pt(0); pj.paragraph_format.space_after = Pt(3)
            rj = pj.add_run(jt)
            rj.font.size = Pt(11); rj.font.color.rgb = DGRAY; rj.font.name = FONT; rj.italic = True

    cl = parsed["sections"].get("contact", [])
    if cl:
        ct = cl[0] if len(cl) == 1 else "  |  ".join(l.strip() for l in cl if l.strip())
        pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_after = Pt(6)
        rc2 = pc.add_run(ct); rc2.font.size = Pt(9.5); rc2.font.color.rgb = DGRAY; rc2.font.name = FONT

    for key in SECTION_ORDER:
        lines = parsed["sections"].get(key, [])
        if not lines: continue
        _section_heading(doc, labels[key])
        for line in lines:
            if line.strip(): _body_line(doc, line)

    doc.add_paragraph()
    pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("Enhanced by SUBUL ·  ATS-Optimized  ·  Powered by SUBUL")
    rf.font.size = Pt(7); rf.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
    rf.font.italic = True; rf.font.name = FONT

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ═════════════════════════════════════════════════════════════════════════════
# FORMAT 2 — BASIC / CLASSIC
# ═════════════════════════════════════════════════════════════════════════════

def generate_docx_basic(parsed: dict) -> bytes:
    F   = "Calibri"; BK  = RGBColor(0x0D,0x0D,0x0D); MGR = RGBColor(0x55,0x55,0x55)
    doc = Document()
    for s in doc.sections: s.top_margin = s.bottom_margin = Inches(0.7); s.left_margin = s.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = F; doc.styles["Normal"].font.size = Pt(10.5)
    pn = doc.add_paragraph(); pn.paragraph_format.space_before = Pt(0); pn.paragraph_format.space_after = Pt(2)
    r = pn.add_run(parsed["name"]); r.bold = True; r.font.size = Pt(32); r.font.color.rgb = BK; r.font.name = F
    jt = parsed.get("job_title", "")
    if jt:
        pj = doc.add_paragraph(); pj.paragraph_format.space_before = Pt(0); pj.paragraph_format.space_after = Pt(4)
        rj = pj.add_run(jt); rj.font.size = Pt(12); rj.font.color.rgb = MGR; rj.font.name = F; rj.italic = True
    ct = _contact_str(parsed)
    if ct:
        pc = doc.add_paragraph(); pc.paragraph_format.space_before = Pt(0); pc.paragraph_format.space_after = Pt(6)
        rc = pc.add_run(ct); rc.font.size = Pt(9); rc.font.color.rgb = MGR; rc.font.name = F
    ps = doc.add_paragraph(); ps.paragraph_format.space_before = Pt(0); ps.paragraph_format.space_after = Pt(10)
    _add_bottom_border(ps, "888888", "8")
    TOTAL = 9936; LW = int(TOTAL * 0.35); RW = TOTAL - LW
    tbl_obj = _make_no_border_table(doc, [LW, RW])
    lcell, rcell = tbl_obj.rows[0].cells[0], tbl_obj.rows[0].cells[1]
    ltcp = lcell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr"); rs = OxmlElement("w:right")
    rs.set(qn("w:val"), "single"); rs.set(qn("w:sz"), "4"); rs.set(qn("w:space"), "12"); rs.set(qn("w:color"), "BBBBBB")
    tcBdr.append(rs); ltcp.append(tcBdr)

    def _csec(tc_el, title):
        p = OxmlElement("w:p"); pPr = OxmlElement("w:pPr")
        sp = OxmlElement("w:spacing"); sp.set(qn("w:before"), "220"); sp.set(qn("w:after"), "60"); pPr.append(sp)
        bd = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "888888")
        bd.append(bot); pPr.append(bd); p.append(pPr)
        r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr"); rPr.append(OxmlElement("w:b"))
        for tag, val in [("w:sz","19"),("w:szCs","19"),("w:spacing","120")]:
            el = OxmlElement(tag); el.set(qn("w:val"), val); rPr.append(el)
        rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), F); rf.set(qn("w:hAnsi"), F); rPr.append(rf)
        r.append(rPr); t = OxmlElement("w:t"); t.text = title.upper(); r.append(t); p.append(r); tc_el.append(p)

    def _cbody(tc_el, line, tab_twips):
        s = line.strip()
        if not s: return
        ib = s.startswith("- ") or s.startswith("• "); clean = s[2:].strip() if ib else s
        p = OxmlElement("w:p"); pPr = OxmlElement("w:pPr")
        sp = OxmlElement("w:spacing"); sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "28"); pPr.append(sp)
        if ib:
            ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "180"); ind.set(qn("w:hanging"), "120"); pPr.append(ind)
        has_tab = "\t" in clean
        if has_tab:
            tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), str(tab_twips)); tabs.append(tab); pPr.append(tabs)
        p.append(pPr)
        def _r(text, bold=False, size=20, col=None):
            r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
            rf2 = OxmlElement("w:rFonts"); rf2.set(qn("w:ascii"), F); rf2.set(qn("w:hAnsi"), F); rPr.append(rf2)
            if bold: rPr.append(OxmlElement("w:b"))
            for tag, val in [("w:sz", str(size)), ("w:szCs", str(size))]: el = OxmlElement(tag); el.set(qn("w:val"), val); rPr.append(el)
            if col: cl = OxmlElement("w:color"); cl.set(qn("w:val"), col); rPr.append(cl)
            r.append(rPr); t = OxmlElement("w:t"); t.text = text
            if text.startswith(" ") or text.endswith(" "): t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t); p.append(r)
        if ib: _r("• ", size=20, col="2B2B2B")
        if has_tab:
            left_t, right_t = clean.split("\t", 1)
            for part in re.split(r'(\*\*[^*]+\*\*)', left_t):
                if not part: continue
                b = part.startswith("**") and part.endswith("**")
                _r(part[2:-2] if b else part, bold=b, size=20, col="2B2B2B")
            rt = OxmlElement("w:r"); tt = OxmlElement("w:t"); tt.text = "\t"; rt.append(tt); p.append(rt)
            _r(right_t.strip(), size=18, col="999999")
        else:
            for part in re.split(r'(\*\*[^*]+\*\*)', clean):
                if not part: continue
                b = part.startswith("**") and part.endswith("**")
                _r(part[2:-2] if b else part, bold=b, size=20, col="2B2B2B")
        tc_el.append(p)

    for cell in [lcell, rcell]:
        for p in cell._tc.findall(qn("w:p")): cell._tc.remove(p)
    LT = int(LW * 0.92); RT = int(RW * 0.93)
    for key in ["education","skills","languages","certifications"]:
        lines = parsed["sections"].get(key, [])
        if not lines: continue
        label = "Certifications" if key == "certifications" else SECTION_LABELS[key]
        _csec(lcell._tc, label)
        for line in lines: _cbody(lcell._tc, line, LT)
    for key, label in [("profile","Summary"),("experience","Experience"),("projects","Projects")]:
        lines = parsed["sections"].get(key, [])
        if not lines: continue
        _csec(rcell._tc, label)
        for line in lines: _cbody(rcell._tc, line, RT)
    pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.paragraph_format.space_before = Pt(10)
    rf = pf.add_run("Enhanced by SUBUL  ·  Basic Classic Format")
    rf.font.size = Pt(7); rf.font.color.rgb = RGBColor(0xBB,0xBB,0xBB); rf.font.italic = True; rf.font.name = F
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()


# ═════════════════════════════════════════════════════════════════════════════
# FORMAT 3 — MODERN (navy blue sidebar)
# ═════════════════════════════════════════════════════════════════════════════

def generate_docx_modern(parsed: dict) -> bytes:
    F = "Calibri"
    WH = RGBColor(0xFF,0xFF,0xFF); BK = RGBColor(0x11,0x11,0x11)
    DGR = RGBColor(0x2B,0x2B,0x2B); LGR = RGBColor(0x88,0x88,0x88)
    PALE = RGBColor(0xBB,0xCC,0xEE)
    doc = Document()
    for s in doc.sections: s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0)
    doc.styles["Normal"].font.name = F; doc.styles["Normal"].font.size = Pt(10)
    TOTAL = 11906; LW = int(TOTAL * 0.33); RW = TOTAL - LW; PAGE_H = 16838
    table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"; tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tw = OxmlElement("w:tblW"); tw.set(qn("w:w"), str(TOTAL)); tw.set(qn("w:type"), "dxa"); tblPr.append(tw)
    bd = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}"); b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0"); b.set(qn("w:space"), "0"); b.set(qn("w:color"), "auto"); bd.append(b)
    tblPr.append(bd)
    grid = OxmlElement("w:tblGrid")
    for w in [LW, RW]: gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    tbl.insert(0, grid)
    row = table.rows[0]; trPr = OxmlElement("w:trPr"); trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(PAGE_H)); trH.set(qn("w:hRule"), "atLeast"); trPr.append(trH); row._tr.insert(0, trPr)
    lcell = table.rows[0].cells[0]; rcell = table.rows[0].cells[1]
    for cell, w in [(lcell,LW),(rcell,RW)]:
        tcp = cell._tc.get_or_add_tcPr(); cw = OxmlElement("w:tcW"); cw.set(qn("w:w"), str(w)); cw.set(qn("w:type"), "dxa"); tcp.append(cw)
    ltcp = lcell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "1A3C72"); ltcp.append(shd)
    rtcp = rcell._tc.get_or_add_tcPr(); shd2 = OxmlElement("w:shd"); shd2.set(qn("w:val"), "clear"); shd2.set(qn("w:color"), "auto"); shd2.set(qn("w:fill"), "FFFFFF"); rtcp.append(shd2)
    PL = Inches(0.28); PR = Inches(0.38)

    def l_sec(title):
        p = lcell.add_paragraph(); p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6); p.paragraph_format.left_indent = PL
        _add_bottom_border(p, "FFFFFF", "4"); r = p.add_run(title.upper()); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WH; r.font.name = F
        rPr = r._r.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), "60"); rPr.append(sp)
    def l_contact(icon, text):
        p = lcell.add_paragraph(); p.paragraph_format.space_after = Pt(5); p.paragraph_format.left_indent = PL
        ri = p.add_run(icon + "  "); ri.font.size = Pt(9); ri.font.color.rgb = PALE; ri.font.name = F
        rv = p.add_run(text); rv.font.size = Pt(9); rv.font.color.rgb = WH; rv.font.name = F
    def l_body(line):
        s = line.strip()
        if not s: return
        ib = s.startswith("- ") or s.startswith("• "); clean = s[2:].strip() if ib else s
        p = lcell.add_paragraph(); p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = PL
        if ib:
            p.paragraph_format.left_indent = Inches(PL.inches + 0.1); p.paragraph_format.first_line_indent = Inches(-0.1)
            rb = p.add_run("•  "); rb.font.size = Pt(9.5); rb.font.color.rgb = WH; rb.font.name = F
        _add_runs(p, clean, Pt(9.5), WH, F)
    def r_sec(title):
        p = rcell.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(5); p.paragraph_format.left_indent = PR
        _add_bottom_border(p, "CCCCCC", "4"); r = p.add_run(title.upper()); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BK; r.font.name = F
        rPr = r._r.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), "80"); rPr.append(sp)
    def r_body(line):
        s = line.strip()
        if not s: return
        ib = s.startswith("- ") or s.startswith("• "); clean = s[2:].strip() if ib else s
        p = rcell.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.left_indent = PR
        if ib:
            p.paragraph_format.left_indent = Inches(PR.inches + 0.18); p.paragraph_format.first_line_indent = Inches(-0.15)
            rb = p.add_run("•  "); rb.font.size = Pt(10); rb.font.color.rgb = DGR; rb.font.name = F
        if "\t" in clean:
            _set_right_tab(p, 5.3); lt, rt = clean.split("\t", 1)
            _add_runs(p, lt, Pt(10), DGR, F); p.add_run("\t").font.name = F
            rdt = p.add_run(rt.strip()); rdt.font.size = Pt(9); rdt.font.color.rgb = LGR; rdt.font.name = F
        else:
            _add_runs(p, clean, Pt(10), DGR, F)

    lcell.paragraphs[0]._p.getparent().remove(lcell.paragraphs[0]._p)
    pnm = lcell.add_paragraph(); pnm.paragraph_format.space_before = Pt(32); pnm.paragraph_format.space_after = Pt(2); pnm.paragraph_format.left_indent = PL
    rnm = pnm.add_run(parsed["name"].upper()); rnm.bold = True; rnm.font.size = Pt(22); rnm.font.color.rgb = WH; rnm.font.name = F
    jt = parsed.get("job_title", "")
    if jt:
        pjt = lcell.add_paragraph(); pjt.paragraph_format.space_before = Pt(0); pjt.paragraph_format.space_after = Pt(10); pjt.paragraph_format.left_indent = PL
        rjt = pjt.add_run(jt); rjt.font.size = Pt(9.5); rjt.font.color.rgb = PALE; rjt.font.name = F; rjt.italic = True
    l_sec("Contact")
    ct = _contact_str(parsed); icons = ["-", "-", "-", "in", "-", "-"]
    for i, item in enumerate(v.strip() for v in re.split(r"\s*\|\s*", ct) if v.strip()):
        l_contact(icons[i] if i < len(icons) else "-", item)
    if parsed["sections"].get("profile"):        l_sec("Personal Statement"); [l_body(l) for l in parsed["sections"]["profile"]]
    if parsed["sections"].get("skills"):         l_sec("Key Skills");         [l_body(l) for l in parsed["sections"]["skills"]]
    if parsed["sections"].get("languages"):      l_sec("Languages");          [l_body(l) for l in parsed["sections"]["languages"]]
    if parsed["sections"].get("certifications"): l_sec("Certifications");     [l_body(l) for l in parsed["sections"]["certifications"]]
    rcell.paragraphs[0]._p.getparent().remove(rcell.paragraphs[0]._p)
    pad = rcell.add_paragraph(); pad.paragraph_format.space_before = Pt(32)
    for key in ["experience","education","projects"]:
        lines = parsed["sections"].get(key, [])
        if not lines: continue
        r_sec(SECTION_LABELS[key]); [r_body(l) for l in lines]
    pf = rcell.add_paragraph(); pf.paragraph_format.space_before = Pt(24); pf.paragraph_format.left_indent = PR
    rf = pf.add_run("Enhanced by SUBUL  ·  Modern Format  ·  Powered by SUBUL")
    rf.font.size = Pt(7); rf.font.color.rgb = RGBColor(0xBB,0xBB,0xBB); rf.font.italic = True; rf.font.name = F
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()


# ═════════════════════════════════════════════════════════════════════════════
# FORMAT 4 — CANADIAN
# ═════════════════════════════════════════════════════════════════════════════

def generate_docx_canadian(parsed: dict) -> bytes:
    F = "Calibri"; NV = RGBColor(0x1F,0x49,0x7D); GR = RGBColor(0x33,0x33,0x33); LG = RGBColor(0x55,0x55,0x55)
    doc = Document()
    for s in doc.sections: s.top_margin = s.bottom_margin = Inches(0.75); s.left_margin = s.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = F; doc.styles["Normal"].font.size = Pt(11)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(parsed["name"].upper()); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NV; r.font.name = F
    ct = _contact_str(parsed)
    if ct:
        p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(ct); r2.font.size = Pt(10); r2.font.color.rgb = LG; r2.font.name = F
    ps = doc.add_paragraph(); ps.paragraph_format.space_after = Pt(4); _add_bottom_border(ps, "1F497D", "12")
    for key in SECTION_ORDER:
        lines = parsed["sections"].get(key, [])
        if not lines: continue
        ph = doc.add_paragraph(); ph.paragraph_format.space_before = Pt(10); ph.paragraph_format.space_after = Pt(3)
        _add_bottom_border(ph, "1F497D", "6")
        rh = ph.add_run(SECTION_LABELS[key].upper()); rh.bold = True; rh.font.size = Pt(11); rh.font.color.rgb = NV; rh.font.name = F
        for line in lines: _render_line(doc, line, F, Pt(11), GR, NV, 6.9)
    pn2 = doc.add_paragraph(); pn2.paragraph_format.space_before = Pt(10)
    rn2 = pn2.add_run("✓ Canadian standard: no photo · no age · no gender · no marital status")
    rn2.font.size = Pt(8); rn2.font.italic = True; rn2.font.color.rgb = RGBColor(0x99,0x99,0x99); rn2.font.name = F
    pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("Enhanced by SUBUL  ·  Canadian Format")
    rf.font.size = Pt(7); rf.font.color.rgb = RGBColor(0xAA,0xAA,0xAA); rf.font.italic = True; rf.font.name = F
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()


# ═════════════════════════════════════════════════════════════════════════════
# FORMAT 5 — EUROPASS
# ═════════════════════════════════════════════════════════════════════════════

def generate_docx_europass(parsed: dict) -> bytes:
    F = "Arial"; EU = RGBColor(0x00,0x33,0x99); GR = RGBColor(0x44,0x44,0x44); WH = RGBColor(0xFF,0xFF,0xFF)
    doc = Document()
    for s in doc.sections: s.top_margin = s.bottom_margin = Inches(0.75); s.left_margin = s.right_margin = Inches(0.9)
    doc.styles["Normal"].font.name = F; doc.styles["Normal"].font.size = Pt(10)
    def euro_band(title):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4); _shading(p, "003399")
        r = p.add_run(f"  {title.upper()}  "); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WH; r.font.name = F
    def euro_field(label, value):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.left_indent = Inches(0.15)
        rl = p.add_run(f"{label}:  "); rl.bold = True; rl.font.size = Pt(9.5); rl.font.color.rgb = EU; rl.font.name = F
        rv = p.add_run(value); rv.font.size = Pt(9.5); rv.font.color.rgb = GR; rv.font.name = F
    pe = doc.add_paragraph(); pe.paragraph_format.space_before = Pt(0); pe.paragraph_format.space_after = Pt(0); _shading(pe, "003399")
    re2 = pe.add_run("  ★  EUROPASS CURRICULUM VITAE"); re2.bold = True; re2.font.size = Pt(9); re2.font.color.rgb = WH; re2.font.name = F
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    pname = doc.add_paragraph(); pname.paragraph_format.space_after = Pt(2)
    rn = pname.add_run(parsed["name"]); rn.bold = True; rn.font.size = Pt(18); rn.font.color.rgb = EU; rn.font.name = F
    euro_band("Personal Information")
    ct = _contact_str(parsed); fl = ["Email","Phone","Address","LinkedIn","GitHub","Portfolio"]
    for i, val in enumerate(re.split(r"\s*\|\s*", ct)):
        val = val.strip()
        if val: euro_field(fl[i] if i < len(fl) else "Other", val)
    euro_map = {
        "profile": "Professional Profile", "experience": "Work Experience",
        "education": "Education and Training", "projects": "Projects",
        "skills": "Skills and Competences", "certifications": "Certificates and Diplomas",
        "languages": "Language Skills (CEFR)",
    }
    for key, label in euro_map.items():
        lines = parsed["sections"].get(key, [])
        if not lines: continue
        euro_band(label)
        for line in lines: _render_line(doc, line, F, Pt(10), GR, EU, 6.9)
    pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("Enhanced by SUBUL  ·  Europass Format  ·  Powered by SUBUL")
    rf.font.size = Pt(7); rf.font.color.rgb = RGBColor(0xAA,0xAA,0xAA); rf.font.italic = True; rf.font.name = F
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()


# ═════════════════════════════════════════════════════════════════════════════
# FORMAT ROUTER
# ═════════════════════════════════════════════════════════════════════════════

def generate_docx(parsed: dict, cv_format: str, photo_bytes=None, language: str = "English") -> bytes:
    if cv_format == "basic":    return generate_docx_basic(parsed)
    if cv_format == "modern":   return generate_docx_modern(parsed)
    if cv_format == "canadian": return generate_docx_canadian(parsed)
    if cv_format == "europass": return generate_docx_europass(parsed)
    return generate_docx_from_llm(parsed, photo_bytes=photo_bytes, language=language)


# ═════════════════════════════════════════════════════════════════════════════
# ENDPOINT 1 — POST /boost-cv
# ═════════════════════════════════════════════════════════════════════════════
@enhance_router.get("/job-data/{job_id:path}")
async def get_job_data(job_id: str):
    """Fetch job details from CosmosDB by URL or id_job."""
    from urllib.parse import unquote
    from database import _get_containers
    import asyncio

    job_id_decoded = unquote(job_id)

    try:
        _, container_jobs, _ = _get_containers()

        # Try by URL first (job_id is the job URL)
        rows = list(await asyncio.to_thread(lambda: list(container_jobs.query_items(
            query="SELECT * FROM c WHERE c.url=@url OFFSET 0 LIMIT 1",
            parameters=[{"name": "@url", "value": job_id_decoded}],
            enable_cross_partition_query=True
        ))))

        # Fallback: try by id_job
        if not rows:
            rows = list(await asyncio.to_thread(lambda: list(container_jobs.query_items(
                query="SELECT * FROM c WHERE c.id_job=@id OFFSET 0 LIMIT 1",
                parameters=[{"name": "@id", "value": job_id_decoded}],
                enable_cross_partition_query=True
            ))))

        if not rows:
            return JSONResponse({"error": "Job not found"}, status_code=404)

        job = rows[0]
        return JSONResponse({
            "title":        job.get("title", ""),
            "description":  job.get("description", ""),
            "requirements": job.get("must_have", ""),
            "location":     job.get("location", ""),
            "seniority":    job.get("seniority", ""),
            "source":       job.get("source", ""),
            "url":          job.get("url", ""),
        })
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
@enhance_router.post("/boost-cv")
async def boost_cv(
    file:             UploadFile = File(...),
    cv_format:        str = Form("ats"),
    include_quiz:     str = Form("true"),
    include_labs:     str = Form("[]"),
    include_certs:    str = Form("[]"),
    extra_data:       str = Form("{}"),
    skipped_sections: str = Form("[]"),
    user_id:          str = Form(""),
    labs_data:        str = Form("[]"),
    certs_data:       str = Form("[]"),
    quiz_data:        str = Form("null"),
):
    content  = await file.read()
    filename = file.filename.lower()
    log.info("📥 Received file: %s (%d bytes) | format: %s", filename, len(content), cv_format)

    # ── Parse skipped sections ────────────────────────────────────────────────
    try:
        skipped = json.loads(skipped_sections)
        if not isinstance(skipped, list):
            skipped = []
    except Exception:
        skipped = []
    log.info("🚫 Skipped sections: %s", skipped)

    # ── Fetch REAL platform data from PostgreSQL ──────────────────────────────
    platform_data = get_platform_data_or_fallback(user_id)
    log.info("✓ Platform data for user '%s': status=%s, %d labs, %d certs",
             user_id, platform_data.get("status"), len(platform_data["labs"]), len(platform_data["certifications"]))
    _no_db_data = platform_data.get("status") in ("no_data","db_error","no_user_id","invalid_user")
    if _no_db_data or (not platform_data["labs"] and not platform_data["certifications"]):
        try:
            _lr = json.loads(labs_data)  if labs_data  != "[]"   else []
            _cr = json.loads(certs_data) if certs_data != "[]"   else []
            _qr = json.loads(quiz_data)  if quiz_data  != "null" else None
            if _lr or _cr or _qr:
                platform_data = {"status":"ok","labs":_lr,"certifications":_cr,"quiz":_qr}
                log.info("✓ Using frontend-supplied platform data: %d labs, %d certs", len(_lr), len(_cr))
        except Exception as _e:
            log.warning("Could not parse frontend platform data: %s", _e)

    # ── 1. Extract text + photo ───────────────────────────────────────────────
    photo_bytes = None
    if filename.endswith(".pdf"):
        raw_text    = extract_text_from_pdf(content)
        photo_bytes = extract_photo_from_pdf(content)
    elif filename.endswith(".docx"):
        raw_text    = extract_text_from_docx(content)
        photo_bytes = extract_photo_from_docx(content)
    elif filename.endswith(".txt"):
        raw_text = content.decode("utf-8", errors="ignore")
    else:
        return JSONResponse(
            {"error": "Unsupported file type. Please upload PDF, DOCX, or TXT."},
            status_code=400
        )

    log.info("📄 Extracted text: %d chars", len(raw_text))

    # ── 2. Detect domain + generate semantic keywords ─────────────────────────
    domain          = detect_domain(raw_text)
    domain_keywords = get_domain_keywords(domain)

    # ── 3. ATS score BEFORE ───────────────────────────────────────────────────
    cv_sections_before = parse_cv_sections(raw_text)
    score_before       = calculate_ats_score(cv_sections_before, raw_text, domain_keywords=domain_keywords)
    log.info("📊 ATS score BEFORE: %d/100", score_before["total"])

    # ── 3b. xAI explanation ───────────────────────────────────────────────────
    explanation_before = explain_ats_score(raw_text, score_before, domain)

    # ── 3c. Detect truly missing sections ────────────────────────────────────
    missing_sections = []
    if not cv_sections_before.get("experience"):
        missing_sections.append("experience")
    if not cv_sections_before.get("education"):
        missing_sections.append("education")
    if not cv_sections_before.get("languages"):
        missing_sections.append("languages")
    log.info("🔍 Missing sections: %s", missing_sections)

    # ── 4. Guard against empty/unfilled templates ─────────────────────────────
    clean_text = _clean_noise(raw_text)
    ok, reason = _has_real_content(clean_text)
    if not ok:
        return JSONResponse({"error": reason}, status_code=422)

    # ── 4b. Inject extra data from modal ─────────────────────────────────────
    try:
        extra = json.loads(extra_data) if extra_data and extra_data.strip() not in ("{}", "") else {}
    except Exception:
        extra = {}

    extra_blocks = []
    if extra.get("languages") and "languages" not in skipped:
        lines = ["LANGUAGES"]
        for row in extra["languages"]:
            lang  = str(row.get("language", "")).strip()
            level = str(row.get("level",    "")).strip()
            if lang:
                lines.append(f"- {lang}: {level}" if level else f"- {lang}")
        if len(lines) > 1:
            extra_blocks.append("# USER-PROVIDED DATA — INCLUDE IN CV:\n" + "\n".join(lines))
            log.info("✓ Injecting languages: %s", lines[1:])

    if extra.get("education") and "education" not in skipped:
        lines = ["EDUCATION"]
        for row in extra["education"]:
            degree     = str(row.get("degree",     "")).strip()
            university = str(row.get("university", "")).strip()
            start      = str(row.get("start",      "")).strip()
            end        = str(row.get("end",         "")).strip()
            date_str   = f"{start} – {end}" if start and end else start or end
            parts = [p for p in [degree, university, date_str] if p]
            if parts:
                lines.append("  " + " | ".join(parts))
        if len(lines) > 1:
            extra_blocks.append("\n".join(lines))
            log.info("✓ Injecting education: %s", lines[1:])

    if extra.get("experience") and "experience" not in skipped:
        lines = ["PROFESSIONAL EXPERIENCE"]
        for row in extra["experience"]:
            title    = str(row.get("title",       "")).strip()
            company  = str(row.get("company",     "")).strip()
            location = str(row.get("location",    "")).strip()
            start    = str(row.get("start",       "")).strip()
            end      = str(row.get("end",         "")).strip()
            desc     = str(row.get("description", "")).strip()
            date_str = f"{start} – {end}" if start and end else start or end
            header_parts = [p for p in [title, company, location] if p]
            if header_parts:
                header = " | ".join(header_parts)
                lines.append(f"**{header}**\t{date_str}" if date_str else f"**{header}**")
                if desc:
                    for sentence in desc.split("."):
                        s = sentence.strip()
                        if s:
                            lines.append(f"- {s}")
        if len(lines) > 1:
            extra_blocks.append("\n".join(lines))
            log.info("✓ Injecting experience: %d entries", len(extra["experience"]))

    if extra_blocks:
        clean_text = clean_text + "\n\n" + "\n\n".join(extra_blocks)

    # ── 5. LLM: generate enhanced CV ─────────────────────────────────────────
    try:
        lab_ids  = json.loads(include_labs)
        cert_ids = json.loads(include_certs)
        quiz_ok  = include_quiz.lower() == "true"

        clean_text = fix_garbled_certifications(clean_text)
        llm_text = generate_cv_with_llm(
            clean_text, platform_data,
            include_quiz=quiz_ok,
            include_lab_ids=lab_ids,
            include_cert_ids=cert_ids,
            skipped_sections=skipped,
        )
        log.info("✓ LLM output: %d chars", len(llm_text))

        if llm_text.strip().startswith("ERROR_TEMPLATE"):
            return JSONResponse({
                "error": (
                    "Your CV appears to be an unfilled template. "
                    "Please replace all placeholder text with your real information and re-upload."
                )
            }, status_code=422)

        if len(llm_text) < 100:
            return JSONResponse({"error": f"LLM response too short: {repr(llm_text)}"}, status_code=500)

    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse({"error": f"LLM error: {str(e)}"}, status_code=500)

    # ── 6. Parse LLM output + ATS score AFTER ────────────────────────────────
    parsed_llm = parse_llm_output(llm_text, skipped_sections=skipped)
    _stated_title = ""
    _sec_kw = re.compile(r"^(profile|résumé|summary|contact|education|expérience|skills"
                         r"|compétences|formation|certif|languages|langues|projects|projets)", re.I)
    for _cand in [l.strip() for l in raw_text.split("\n") if l.strip()][1:5]:
        if "@" in _cand or re.search(r"\+?\d[\d\s\-]{6,}", _cand): continue
        if _sec_kw.match(_cand) or _cand.startswith(("•","-","·")): continue
        if 3 <= len(_cand) <= 100 and re.search(r"[a-zA-ZÀ-ÿ]", _cand):
            _stated_title = _cand
            log.info("✓ Stated job title from raw CV: %s", _stated_title)
            break
    parsed_llm["job_title"] = infer_job_title(
        platform_data, parsed_llm["sections"],
        cv_domain=domain, cv_stated_title=_stated_title,
    )
    score_after = calculate_ats_score(
        {**parsed_llm["sections"], "header": [parsed_llm["name"]] + parsed_llm["sections"].get("contact", [])},
        llm_text,
        domain_keywords=domain_keywords,
    )
    log.info("📊 ATS score AFTER: %d/100 (Δ %+d)", score_after["total"], score_after["total"] - score_before["total"])
    # ── 6b. Extract structured data + save to CosmosDB ───────────────────────
    log.info("🔑 user_id reçu: '%s'", user_id)
    if user_id:
        try:
            cv_structured = extract_cv_structured(llm_text)
            cv_structured["domain"] = domain
            cv_structured["level"]  = (platform_data.get("quiz") or {}).get("level", "")
            cv_structured["ats_score_before"] = score_before["total"]
            cv_structured["ats_score_after"] = score_after["total"]
            save_cv_to_cosmos(user_id, cv_structured)

    
        except Exception as e:
            log.error("⚠ CosmosDB step skipped: %s", str(e))
    else:
        log.warning("⚠ user_id vide — CosmosDB ignoré")
    # ── 7. Generate DOCX ──────────────────────────────────────────────────────
    # Detect language from raw_text for correct section label language
    _fr_detect = ["et", "ou", "les", "des", "dans", "pour", "avec", "sur",
                  "une", "est", "par", "expérience", "formation", "compétences", "langues"]
    _raw_lo   = raw_text.lower()
    _fr_hits  = sum(1 for w in _fr_detect if f" {w} " in _raw_lo)
    _doc_lang = "French" if _fr_hits >= 4 else "English"
    log.info("📝 Generating DOCX — format: %s | lang: %s | sections: %s",
             cv_format, _doc_lang, list(parsed_llm["sections"].keys()))
    docx_bytes = generate_docx(parsed_llm, cv_format, photo_bytes=photo_bytes, language=_doc_lang)

    # ── 8. Return ─────────────────────────────────────────────────────────────
    name = parsed_llm["name"].split()[0] if parsed_llm["name"] else "CV"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition":          f'attachment; filename="{name}_CV_{cv_format.upper()}.docx"',
            "X-ATS-Score-Before":           str(score_before["total"]),
            "X-ATS-Score-After":            str(score_after["total"]),
            "X-ATS-Breakdown-Before":       json.dumps(score_before["breakdown"]),
            "X-ATS-Breakdown-After":        json.dumps(score_after["breakdown"]),
            "X-ATS-Explanation":            json.dumps(explanation_before),
            "X-Parsed-CV":                  json.dumps(parsed_llm),
            "X-Domain":                     domain,
            "X-Keywords-Matched":           json.dumps(score_after["breakdown"]["keywords"].get("matched", [])),
            "X-Missing-Sections":           json.dumps(missing_sections),
            "Access-Control-Expose-Headers":
                "X-ATS-Score-Before,X-ATS-Score-After,X-ATS-Breakdown-Before,"
                "X-ATS-Breakdown-After,X-ATS-Explanation,X-Parsed-CV,"
                "X-Domain,X-Keywords-Matched,X-Missing-Sections",
        },
    )


# ═════════════════════════════════════════════════════════════════════════════
# ENDPOINT 2 — POST /apply-format
# ═════════════════════════════════════════════════════════════════════════════

@enhance_router.post("/apply-format")
async def apply_format(
    parsed_cv: str = Form(...),
    cv_format: str = Form("ats"),
):
    try:
        parsed = json.loads(parsed_cv)
    except Exception:
        return JSONResponse({"error": "Invalid parsed_cv JSON."}, status_code=400)

    try:
        docx_bytes = generate_docx(parsed, cv_format)
    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse({"error": f"Format error: {str(e)}"}, status_code=500)

    name = parsed.get("name", "CV").split()[0]
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition":           f'attachment; filename="{name}_CV_{cv_format.upper()}.docx"',
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )


# ═════════════════════════════════════════════════════════════════════════════
# HEALTH CHECK
# ═════════════════════════════════════════════════════════════════════════════

@enhance_router.get("/")
def root():
    return {
        "status":    "CV Booster API v4 — Azure OpenAI",
        "endpoints": ["POST /boost-cv", "POST /apply-format", "GET /platform-data/{user_id}"],
        "fitz":      FITZ_OK,
        "semantic":  SEMANTIC_OK,
    }


# ═════════════════════════════════════════════════════════════════════════════
# ENDPOINT 3 — GET /platform-data/{user_id}
# ═════════════════════════════════════════════════════════════════════════════


@enhance_router.get("/platform-data/{user_id}")
async def get_platform_data(user_id: str = FPath(...)):
    """
    Returns real platform data (quiz + labs + certs + recommendations) for a user.
    Falls back gracefully if user not found or DB unreachable.
    """
    data = get_platform_data_or_fallback(user_id)

    # Add recommendations only if user has valid data
    if data.get("status") == "ok":
        quiz_domain = (data.get("quiz") or {}).get("domain", "")
        recs = fetch_recommendations(user_id, quiz_domain)
        data["recommendations"] = recs
    else:
        data["recommendations"] = {"certifications": [], "labs": []}

    return JSONResponse(data)
# ═════════════════════════════════════════════════════════════════════════════
# ENDPOINT 4 — POST /enhance-cv-for-job
# ═════════════════════════════════════════════════════════════════════════════

@enhance_router.post("/enhance-cv-for-job")
async def enhance_cv_for_job(
    user_id:   str = Form(""),
    job_id:    str = Form(""),
    cv_format: str = Form("ats"),
):
    """
    Enhance the user's CV specifically for a job offer.
    - Loads CV text from CosmosDB (cv_raw_text field)
    - Loads job details from CosmosDB (by URL)
    - Rewrites the CV targeting that specific job
    - Returns a DOCX file
    """
    from urllib.parse import unquote
    from database import get_user, _get_containers
    import asyncio

    job_id = unquote(job_id)
    log.info("📥 /enhance-cv-for-job — user_id=%s job_id=%s", user_id, job_id[:60])

    # ── 1. Load user CV from CosmosDB ─────────────────────────────────────────
    if not user_id:
        return JSONResponse({"error": "user_id is required"}, status_code=400)

    user = await get_user(int(user_id))
    if not user:
        return JSONResponse({"error": f"User {user_id} not found"}, status_code=404)

    cv_raw_text = user.get("cv_raw_text", "").strip()
    if not cv_raw_text or len(cv_raw_text) < 50:
        return JSONResponse({
            "error": "No CV found for this user. Please run a scan first to save your CV."
        }, status_code=400)

    # ── 2. Load job details from CosmosDB ─────────────────────────────────────
    job_title = ""
    job_description = ""
    job_requirements = ""

    if job_id:
        try:
            _, container_jobs, _ = _get_containers()
            rows = list(await asyncio.to_thread(lambda: list(container_jobs.query_items(
                query="SELECT * FROM c WHERE c.url=@url OFFSET 0 LIMIT 1",
                parameters=[{"name": "@url", "value": job_id}],
                enable_cross_partition_query=True
            ))))
            if rows:
                job = rows[0]
                job_title       = job.get("title", "")
                job_description = job.get("description", "")
                job_requirements = job.get("must_have", "") or job.get("requirements", "")
                log.info("✓ Job found: %s", job_title)
            else:
                log.warning("⚠ Job not found in DB for url: %s", job_id[:60])
        except Exception as e:
            log.error("Job fetch error: %s", e)

    # ── 3. Build job-targeted system prompt ───────────────────────────────────
    job_context = ""
    if job_title or job_description or job_requirements:
        job_context = f"""
TARGET JOB OFFER — TAILOR THE CV SPECIFICALLY FOR THIS ROLE:
- Job Title: {job_title}
- Required Skills / Requirements: {job_requirements[:800]}
- Job Description: {job_description[:1200]}

INSTRUCTIONS FOR THIS JOB:
- Reorder and emphasize skills that match this job's requirements
- Use keywords from the job description in the profile and skills sections
- Highlight relevant experience that matches this role
- Keep all factual information — do NOT invent or exaggerate
"""

    # ── 4. Detect domain + platform data ──────────────────────────────────────
    domain          = detect_domain(cv_raw_text)
    domain_keywords = get_domain_keywords(domain)
    platform_data   = get_platform_data_or_fallback(user_id)
    clean_text      = _clean_noise(cv_raw_text)

    # ── 5. Validate CV content ────────────────────────────────────────────────
    ok, reason = _has_real_content(clean_text)
    if not ok:
        return JSONResponse({"error": reason}, status_code=422)

    # ── 6. Generate enhanced CV with job context injected ─────────────────────
    try:
        # Inject job context into the CV text so the LLM sees it
        cv_with_job = clean_text
        if job_context:
            cv_with_job = clean_text + "\n\n" + job_context

        clean_text = fix_garbled_certifications(cv_with_job)

        llm_text = generate_cv_with_llm(
            clean_text,
            platform_data,
            include_quiz=True,
            include_lab_ids=None,
            include_cert_ids=None,
            skipped_sections=[],
        )

        if llm_text.strip().startswith("ERROR_TEMPLATE"):
            return JSONResponse({"error": "CV appears to be an empty template."}, status_code=422)

        if len(llm_text) < 100:
            return JSONResponse({"error": f"LLM response too short: {repr(llm_text)}"}, status_code=500)

    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse({"error": f"LLM error: {str(e)}"}, status_code=500)

    # ── 7. Parse + score ──────────────────────────────────────────────────────
    parsed_llm = parse_llm_output(llm_text, skipped_sections=[])

    # Infer job title
    _sec_kw = re.compile(
        r"^(profile|résumé|summary|contact|education|expérience|skills"
        r"|compétences|formation|certif|languages|langues|projects|projets)", re.I
    )
    _stated_title = job_title or ""
    if not _stated_title:
        for _cand in [l.strip() for l in cv_raw_text.split("\n") if l.strip()][1:5]:
            if "@" in _cand or re.search(r"\+?\d[\d\s\-]{6,}", _cand): continue
            if _sec_kw.match(_cand) or _cand.startswith(("•", "-", "·")): continue
            if 3 <= len(_cand) <= 100 and re.search(r"[a-zA-ZÀ-ÿ]", _cand):
                _stated_title = _cand; break

    parsed_llm["job_title"] = infer_job_title(
        platform_data, parsed_llm["sections"],
        cv_domain=domain, cv_stated_title=_stated_title,
    )

    score_before = calculate_ats_score(
        parse_cv_sections(cv_raw_text), cv_raw_text, domain_keywords=domain_keywords
    )
    score_after = calculate_ats_score(
        {**parsed_llm["sections"], "header": [parsed_llm["name"]] + parsed_llm["sections"].get("contact", [])},
        llm_text, domain_keywords=domain_keywords,
    )
    log.info("📊 Score: %d → %d (+%d)", score_before["total"], score_after["total"],
             score_after["total"] - score_before["total"])

    # ── 8. Detect language + generate DOCX ───────────────────────────────────
    _fr_detect = ["et","ou","les","des","dans","pour","avec","sur","une","est",
                  "par","expérience","formation","compétences","langues"]
    _fr_hits  = sum(1 for w in _fr_detect if f" {w} " in cv_raw_text.lower())
    _doc_lang = "French" if _fr_hits >= 4 else "English"

    docx_bytes = generate_docx(parsed_llm, cv_format, photo_bytes=None, language=_doc_lang)

    # ── 9. Return DOCX ────────────────────────────────────────────────────────
    name = parsed_llm["name"].split()[0] if parsed_llm["name"] else "CV"
    safe_title = re.sub(r"[^\w\s-]", "", job_title).replace(" ", "_")[:30] if job_title else "Job"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition":          f'attachment; filename="{name}_CV_{safe_title}.docx"',
            "X-ATS-Score-Before":           str(score_before["total"]),
            "X-ATS-Score-After":            str(score_after["total"]),
            "X-ATS-Breakdown-Before":       json.dumps(score_before["breakdown"]),
            "X-ATS-Breakdown-After":        json.dumps(score_after["breakdown"]),
            "X-Parsed-CV":                  json.dumps(parsed_llm),
            "X-Domain":                     domain,
            "Access-Control-Expose-Headers":
                "X-ATS-Score-Before,X-ATS-Score-After,X-ATS-Breakdown-Before,"
                "X-ATS-Breakdown-After,X-Parsed-CV,X-Domain",
        },
    )